#!/usr/bin/env python3
"""
Generador Word — Documentación de Usuario Plataforma Algarrobo BASE2
Geoportal Municipal — I. Municipalidad de Algarrobo
Versión 2.0 — Redacción senior técnica expandida
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── COLORES INSTITUCIONALES ──────────────────────────────────────────────────
NAVY      = RGBColor(0x1A, 0x2E, 0x4A)
NAVY_MED  = RGBColor(0x1D, 0x35, 0x57)
YELLOW    = RGBColor(0xF5, 0xB7, 0x00)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT= RGBColor(0xF2, 0xF5, 0xF9)
GRAY_MED  = RGBColor(0xDE, 0xE2, 0xE6)
TEXT_DARK = RGBColor(0x2D, 0x37, 0x48)
TEXT_MED  = RGBColor(0x4A, 0x55, 0x68)
RED_CRIT  = RGBColor(0xC5, 0x30, 0x30)
ORANGE_ALT= RGBColor(0xDD, 0x6B, 0x20)
GREEN_OK  = RGBColor(0x27, 0x6B, 0x3D)
BLUE_INFO = RGBColor(0x2B, 0x6C, 0xB0)
YELLOW_W  = RGBColor(0xB7, 0x79, 0x10)

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      'Documentacion_Algarrobo_BASE2.docx')

# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS DE FORMATO
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, rgb_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        for k, v in kwargs.items():
            tag.set(qn(f'w:{k}'), str(v))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run_bold(para, text, color=None, size=None):
    run = para.add_run(text)
    run.bold = True
    if color: run.font.color.rgb = color
    if size:  run.font.size = Pt(size)
    return run


def cell_para(cell, text, bold=False, color=None, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    para = cell.paragraphs[0]
    para.alignment = align
    set_paragraph_spacing(para)
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return para


def set_paragraph_spacing(para, before=0, after=4, line=1.0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing      = Pt(line * 13)


def heading(doc, text, level=1):
    if level == 1:
        p = doc.add_heading('', level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        pf.left_indent  = Cm(0.4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'),   'single')
        left.set(qn('w:sz'),    '24')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), 'F5B700')
        pBdr.append(left)
        pPr.append(pBdr)
        run = p.add_run(text)
        run.font.color.rgb = NAVY
        run.font.bold = True
        run.font.size = Pt(14)
        return p
    elif level == 2:
        p = doc.add_heading('', level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after  = Pt(4)
        run = p.add_run(text)
        run.font.color.rgb = NAVY_MED
        run.font.bold = True
        run.font.size = Pt(12)
        return p
    else:
        p = doc.add_heading('', level=3)
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.color.rgb = NAVY_MED
        run.font.bold = True
        run.font.size = Pt(11)
        return p


def body(doc, text, indent=False, italic=False, color=None, bold=False, after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(after)
    pf.left_indent  = Cm(0.4) if indent else Cm(0)
    run = p.add_run(text)
    run.font.size   = Pt(10.5)
    run.font.bold   = bold
    run.italic      = italic
    if color: run.font.color.rgb = color
    else:     run.font.color.rgb = TEXT_DARK
    return p


def bullet(doc, text, level=1, bold_prefix=None, color=None):
    p = doc.add_paragraph(style='List Bullet' if level == 1 else 'List Bullet 2')
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.left_indent  = Cm(0.8 * level)
    if bold_prefix:
        run = p.add_run(bold_prefix + ': ')
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY_MED
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color: run.font.color.rgb = color
    else:     run.font.color.rgb = TEXT_DARK
    return p


def numbered(doc, text, level=1):
    p = doc.add_paragraph(style='List Number' if level == 1 else 'List Number 2')
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    return p


def note_box(doc, text, label='Nota', bcolor='2B6CB0', bgcolor='EBF4FF', tcolor=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bgcolor)
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '12')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), bcolor)
        tcBorders.append(tag)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=2, after=2)
    r1 = p.add_run(f'{label}: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(
        int(bcolor[0:2], 16), int(bcolor[2:4], 16), int(bcolor[4:6], 16)
    )
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = tcolor if tcolor else TEXT_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def divider(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'DEE2E6')
    pBdr.append(bot)
    pPr.append(pBdr)


def page_break(doc):
    doc.add_page_break()


def make_table(doc, headers, rows, col_widths=None, header_bg='1A2E4A'):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        cell_para(cell, h, bold=True, color=WHITE, size=9.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = 'F7F9FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell_para(cell, str(val), size=9.5, color=TEXT_DARK, align=align)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


def role_table(doc):
    headers = ['Rol', 'Nivel', 'Descripción Funcional', 'Acceso Principal']
    rows = [
        ['Admin General',       '10', 'Control total: usuarios, configuración global, auditoría, reportes, migración de datos', 'Todos los módulos sin excepción'],
        ['Admin Proyectos',     '11', 'Gestión operativa: crear, editar y eliminar fichas de proyectos y licitaciones activas',  'SECPLAN, Licitaciones, Geoportal'],
        ['Director de Obras',   '12', 'Supervisión técnica: revisar avances de obra, aprobar etapas, analizar informes técnicos', 'Licitaciones (vista supervisión), Geoportal'],
        ['Profesional Técnico', '13', 'Trabajo en terreno: ingresar avances, adjuntar documentos, registrar hitos y fotos',       'SECPLAN (proyectos asignados), Licitaciones'],
        ['Finanzas',            '14', 'Control presupuestario: gestionar montos, fuentes de financiamiento y estados de pago',    'SECPLAN (vista financiera), Licitaciones (estados de pago)'],
        ['Supervisor Alcaldía', '15', 'Visión estratégica ejecutiva: dashboards, reportes agregados, aprobación de proyectos',    'Solo lectura ejecutiva en todos los módulos'],
        ['Transparencia',       '16', 'Acceso para publicación de información ciudadana debidamente autorizada',                  'Solo lectura pública (proyectos publicados)'],
    ]
    return make_table(doc, headers, rows, col_widths=[3.2, 1.3, 7.0, 5.0])


# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('I. MUNICIPALIDAD DE ALGARROBO')
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Unidad de Informática — Secretaría Comunal de Planificación (SECPLAN)')
    r.font.size  = Pt(11)
    r.font.color.rgb = TEXT_MED

    doc.add_paragraph()

    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '12')
        tag.set(qn('w:space'), '4')
        tag.set(qn('w:color'), 'F5B700')
        pBdr.append(tag)
    pPr.append(pBdr)
    div.paragraph_format.space_before = Pt(10)
    div.paragraph_format.space_after  = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('PLATAFORMA GEOPORTAL MUNICIPAL')
    r.font.size  = Pt(26)
    r.font.bold  = True
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Algarrobo BASE2')
    r.font.size  = Pt(20)
    r.font.bold  = True
    r.font.color.rgb = NAVY_MED

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Documentación Integral del Sistema — Versión 2.0')
    r.font.size  = Pt(13)
    r.font.color.rgb = TEXT_MED

    doc.add_paragraph()
    doc.add_paragraph()

    t = doc.add_table(rows=8, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    items = [
        ('01', 'Descripción General del Portal'),
        ('02', 'Guía de Uso Administrador'),
        ('03', 'Guía de Uso Módulo Proyectos'),
        ('04', 'Guía de Uso Licitaciones'),
        ('05', 'Guía de Uso Geoportal'),
        ('06', 'Guía de Migración del Portal'),
        ('07', 'Resolución de Análisis Informática'),
    ]
    for i, (num, name) in enumerate(items):
        c0, c1 = t.rows[i].cells[0], t.rows[i].cells[1]
        set_cell_bg(c0, '1A2E4A')
        set_cell_bg(c1, 'F7F9FC' if i % 2 == 0 else 'EAEFF6')
        cell_para(c0, num,  bold=True, color=WHITE, size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(c1, name, bold=False, color=NAVY,  size=11)
    for col_idx, width in enumerate([1.5, 12.0]):
        for cell in t.columns[col_idx].cells:
            cell.width = Cm(width)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Mayo 2026  —  Unidad de Informática Municipal')
    r.font.size  = Pt(11)
    r.font.color.rgb = TEXT_MED

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN GENERAL DEL PORTAL
# ═══════════════════════════════════════════════════════════════════════════════

def section_1(doc):
    heading(doc, '1. Descripción General del Portal')

    body(doc,
         'El Geoportal Municipal Algarrobo BASE2 es la plataforma digital integrada de gestión '
         'pública de la Secretaría Comunal de Planificación (SECPLAN) de la Ilustre Municipalidad '
         'de Algarrobo. Centraliza en un único entorno web el portafolio completo de inversión '
         'pública comunal, el control detallado de los procesos licitatorios, la auditoría '
         'automatizada de calidad técnica, la inteligencia territorial mediante cartografía '
         'interactiva, y la participación ciudadana a través de reportes geolocalizados.')

    body(doc,
         'La plataforma nace de la necesidad institucional de reemplazar el trabajo en planillas '
         'Excel desconectadas y carpetas de red sin control de versiones, por un sistema de '
         'registro único, trazable y auditable. Está concebida como el Sistema Centralizado de '
         'Control de la Inversión Municipal, transformando datos dispersos en activos estratégicos '
         'disponibles en tiempo real para la toma de decisiones de autoridades, directivos y '
         'equipos técnicos de la municipalidad.')

    heading(doc, '1.1. Contexto Institucional y Problema que Resuelve', level=2)

    body(doc,
         'La SECPLAN de la Municipalidad de Algarrobo es la unidad encargada de planificar, '
         'formular, evaluar y controlar los proyectos de inversión pública comunal. Antes de '
         'la implementación del Geoportal, el seguimiento de proyectos se realizaba mediante '
         'archivos Excel individuales, sin vinculación entre sí, sin control de acceso por rol '
         'ni trazabilidad de cambios. Los procesos licitatorios se controlaban con listas de '
         'verificación en papel. La información geográfica de los proyectos no existía de forma '
         'sistematizada. El Geoportal BASE2 resuelve estas brechas de manera integral.')

    body(doc, 'Principales problemas resueltos por la plataforma:')
    bullet(doc, 'Dispersión de información: todos los proyectos en un único repositorio centralizado con búsqueda y filtros avanzados.')
    bullet(doc, 'Falta de trazabilidad: cada cambio en datos queda registrado con usuario, fecha, hora y contenido anterior y nuevo.')
    bullet(doc, 'Control manual de licitaciones: el workflow de 32 pasos estandarizado guía al equipo paso a paso con alertas automáticas de plazos.')
    bullet(doc, 'Sin análisis de calidad: el motor de auditoría evalúa automáticamente 7 dimensiones con 15 criterios, generando reportes PDF con plan correctivo.')
    bullet(doc, 'Sin inteligencia territorial: el Geoportal visualiza todos los proyectos en un mapa interactivo con capas temáticas, filtros y fichas de detalle.')
    bullet(doc, 'Sin participación ciudadana: la app móvil permite a vecinos reportar problemas en la vía pública con geolocalización y seguimiento de estado.')

    heading(doc, '1.2. Propósito y Alcance del Sistema', level=2)

    body(doc, 'La plataforma cumple cinco funciones estratégicas principales:')
    bullet(doc, 'Registro, seguimiento y control de fichas de proyectos municipales con más de 60 campos estandarizados, organizados en 6 secciones funcionales.',
           bold_prefix='Gestión de Proyectos')
    bullet(doc, 'Control integral del proceso de licitación pública en 32 pasos estandarizados, desde el Acuerdo de Concejo hasta la Recepción Definitiva de obras.',
           bold_prefix='Control Licitatorio')
    bullet(doc, 'Motor automatizado que evalúa proyectos en 7 dimensiones con 15 validaciones específicas (V001-V015), generando reportes PDF individuales con plan de acción correctivo priorizado.',
           bold_prefix='Auditoría de Calidad')
    bullet(doc, 'Mapa interactivo con capas GIS que georreferencia proyectos, reportes ciudadanos y capas temáticas; permite análisis territorial diferenciado por rol de usuario.',
           bold_prefix='Inteligencia Territorial')
    bullet(doc, 'Aplicación web móvil que permite a vecinos reportar problemas en la vía pública (baches, luminarias, basura) con geolocalización GPS y seguimiento de estado de atención.',
           bold_prefix='Participación Ciudadana')

    heading(doc, '1.3. Módulos Funcionales del Sistema', level=2)

    body(doc,
         'La plataforma está compuesta por siete módulos funcionales principales, '
         'cada uno con su propio conjunto de funcionalidades y control de acceso por rol:')

    headers = ['Módulo', 'Descripción Funcional', 'Usuarios Principales', 'Acceso']
    rows = [
        ['SECPLAN / Proyectos',
         'Gestión integral de fichas de proyectos de inversión pública: registro, seguimiento, documentos, hitos, observaciones, geomapas y chat IA por proyecto',
         'Admin General, Admin Proyectos, Técnico, Director',
         'Roles 10-14'],
        ['Licitaciones',
         'Control del proceso licitatorio en 32 pasos estandarizados con seguimiento de fechas, documentos adjuntos y alertas de vencimiento de plazos',
         'Admin General, Admin Proyectos, Director de Obras',
         'Roles 10-12'],
        ['Geoportal (Mapa)',
         'Visualización territorial interactiva de proyectos y reportes ciudadanos con capas temáticas, filtros combinados y fichas de detalle',
         'Todos los roles internos',
         'Roles 10-16'],
        ['Auditoría Integral',
         'Motor de análisis automático de calidad en 7 dimensiones con 15 validaciones, generación de reportes PDF y plan de acción correctivo por proyecto',
         'Admin General exclusivamente',
         'Rol 10'],
        ['Control de Actividad',
         'Trazabilidad completa de acciones de todos los usuarios: login, cambios de datos, subida de archivos, con filtros por usuario, fecha y tipo de acción',
         'Admin General exclusivamente',
         'Rol 10'],
        ['Administración',
         'Gestión de usuarios y roles, mantenimiento de catálogos (áreas, etapas, estados, financiamientos, sectores, lineamientos), gestión de divisiones organizacionales',
         'Admin General exclusivamente',
         'Rol 10'],
        ['App Móvil / Reportes Ciudadanos',
         'Aplicación web adaptada a dispositivos móviles para reporte ciudadano geolocalizados y fiscalización técnica en terreno con captura de fotografías',
         'Vecinos registrados, Fiscalizadores',
         'Acceso público supervisado'],
    ]
    make_table(doc, headers, rows, col_widths=[3.2, 7.0, 3.8, 2.5])

    heading(doc, '1.4. Arquitectura del Sistema', level=2)

    body(doc,
         'El sistema está desplegado sobre una infraestructura de cuatro capas segregadas '
         'que garantizan seguridad, disponibilidad y separación de responsabilidades:')

    headers = ['Capa', 'Tecnología', 'Función', 'Ubicación']
    rows = [
        ['Capa 1 — Hypervisor',
         'Proxmox VE',
         'Virtualización del servidor físico municipal. Gestiona las VMs y contenedores (CT) de forma independiente',
         'Servidor físico on-premise'],
        ['Capa 2 — Proxy Externo',
         'Nginx en LXC Container',
         'Proxy inverso que recibe el tráfico HTTPS externo y lo redirige al contenedor de aplicación. Termina TLS.',
         'CT en Proxmox'],
        ['Capa 3 — Aplicación Docker',
         'Docker + Portainer en VM',
         'Contenedores de backend (Flask/Gunicorn) y frontend (Nginx) orquestados mediante stacks de Portainer',
         'VM en Proxmox'],
        ['Capa 4 — Base de Datos',
         'PostgreSQL en LXC Container',
         'Instancia PostgreSQL dedicada, aislada de la capa de aplicación, con acceso solo desde la red interna del hypervisor',
         'CT en Proxmox'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 3.0, 6.5, 3.5])

    body(doc,
         'El acceso externo se gestiona mediante un Túnel Cloudflare (cloudflared) que establece '
         'un canal seguro de salida sin necesidad de abrir puertos en el firewall del servidor. '
         'Esto elimina la exposición directa de la IP pública municipal.')

    heading(doc, '1.5. Roles de Usuario y Modelo de Permisos (RBAC)', level=2)

    body(doc,
         'El sistema implementa un modelo de control de acceso basado en roles (RBAC) '
         'con siete niveles jerárquicos numerados del 10 al 16. El número de nivel determina '
         'el rango de privilegios: a menor número, mayor acceso. Cada rol accede exclusivamente '
         'a las funcionalidades correspondientes a su responsabilidad institucional. '
         'La validación de permisos se realiza tanto en el navegador del usuario (frontend) '
         'como en cada endpoint de la API del servidor (backend), garantizando que un usuario '
         'no pueda acceder a datos fuera de su nivel aunque conozca la URL directa.')

    role_table(doc)

    note_box(doc,
             'El sistema implementa el principio de menor privilegio: cada usuario solo puede '
             'ver y modificar los datos que corresponden estrictamente a su rol. Un Profesional '
             'Técnico (nivel 13) solo puede editar los proyectos que le han sido asignados como '
             'responsable. El servidor valida el token JWT y el nivel de rol en cada petición API.',
             label='Principio de Seguridad', bcolor='1A2E4A', bgcolor='EEF2F7')

    heading(doc, '1.6. Stack Tecnológico Detallado', level=2)

    body(doc,
         'La plataforma está construida íntegramente con tecnologías de código abierto, '
         'sin dependencias de licencias comerciales. A continuación se describe en detalle '
         'cada componente tecnológico del sistema:')

    headers = ['Componente', 'Tecnología', 'Versión', 'Función Específica en el Sistema']
    rows = [
        ['Frontend Web',
         'HTML5 + Vanilla JavaScript ES6+',
         'Sin framework',
         'Interfaz de usuario sin frameworks pesados. Usa Fetch API para comunicación con el backend. CSS con TailwindCSS para estilos responsivos.'],
        ['Backend API',
         'Python + Flask',
         'Flask 3.0',
         'API REST modular con 9 Blueprints independientes: auth, users, proyectos, licitaciones, documentos, calendario, control, auditoria, mobile. Ejecutado con Gunicorn (4 workers, 2 threads).'],
        ['Base de Datos',
         'PostgreSQL',
         '14+',
         'Motor relacional con soporte PostGIS para datos geográficos. Aloja todas las tablas del sistema: proyectos, licitaciones, usuarios, control de actividad, reportes de auditoría.'],
        ['Autenticación',
         'JWT (JSON Web Tokens)',
         'PyJWT 2.x',
         'Tokens firmados con clave secreta (JWT_SECRET_KEY). Se almacenan en localStorage del navegador y se validan en cada petición API mediante el header Authorization: Bearer <token>.'],
        ['Seguridad de Contraseñas',
         'Bcrypt',
         '4.x',
         'Hashing de contraseñas con salt aleatorio. Las contraseñas nunca se almacenan en texto plano ni son reversibles.'],
        ['Rate Limiting',
         'Flask-Limiter',
         '3.x',
         'Límite de intentos de login (5 por minuto por IP) para prevenir ataques de fuerza bruta.'],
        ['Almacenamiento de Archivos',
         'Volumen Docker persistente /data',
         'Docker Volumes',
         'Tres directorios: /data/docs/ (documentos de proyectos), /data/auditoria_reportes/ (PDF de auditoría), /data/fotos_reportes/ (fotografías ciudadanas).'],
        ['Generación de Reportes PDF',
         'ReportLab',
         '4.4.10',
         'Generación programática de reportes de auditoría en PDF con formato institucional (colores NAVY/YELLOW, tablas de análisis, plan de acción correctivo).'],
        ['OCR y Extracción de Texto',
         'Tesseract + poppler-utils',
         'Tesseract 5.x',
         'Extracción de texto de documentos PDF e imágenes para indexación. Usa el modelo de idioma español (spa).'],
        ['Cartografía Web',
         'Leaflet.js + GeoJSON',
         'Leaflet 1.9',
         'Mapa interactivo con soporte de capas GeoJSON personalizadas. Integración con IDE Chile para capas base del territorio nacional.'],
        ['Asistente IA',
         'GLM-4.5 Flash (ZhipuAI)',
         'API ZhipuAI',
         'Chat asistente por proyecto. Proxy en el backend que requiere sesión activa. La clave API se almacena en variables de entorno del servidor.'],
        ['Proxy de Acceso Externo',
         'Cloudflare Tunnel (cloudflared)',
         '2024+',
         'Canal cifrado de salida sin puertos abiertos. Mapea el dominio geoportal.munialgarrobo.cl al contenedor de aplicación interno.'],
        ['Orquestación de Contenedores',
         'Docker + Portainer',
         'Docker 24+',
         'Docker Compose para orquestación local. Portainer Community Edition para gestión visual de stacks, variables de entorno y logs desde interfaz web.'],
    ]
    make_table(doc, headers, rows, col_widths=[2.8, 2.8, 1.8, 9.1])

    heading(doc, '1.7. Acceso a la Plataforma', level=2)

    body(doc,
         'La plataforma es accesible desde cualquier navegador web moderno a través '
         'del dominio institucional. No requiere instalación de software adicional '
         'en el equipo del usuario.')

    note_box(doc,
             'URL de acceso: https://geoportal.munialgarrobo.cl\n'
             'Las credenciales (usuario y contraseña inicial) son proporcionadas por la '
             'Unidad de Informática de la municipalidad. En el primer acceso se recomienda '
             'cambiar la contraseña desde el perfil de usuario.',
             label='URL Oficial', bcolor='276B3D', bgcolor='EBF5F0')

    body(doc, 'Requisitos del equipo del usuario:')

    headers = ['Requisito', 'Especificación Mínima', 'Recomendado']
    rows = [
        ['Navegador Web', 'Google Chrome 90, Firefox 88 o Edge 90', 'Google Chrome 110+ o Edge 110+ (última versión)'],
        ['Conexión a Internet', 'Banda ancha 5 Mbps simétrico', '10 Mbps o superior para carga fluida de capas GIS y documentos'],
        ['Resolución de Pantalla', '1280 × 720 píxeles', '1920 × 1080 píxeles para visualización óptima del Geoportal'],
        ['JavaScript', 'Habilitado en el navegador', 'Habilitado (requerido, sin él el sistema no funciona)'],
        ['Cookies / localStorage', 'Habilitado en el navegador', 'Requerido para la sesión JWT del usuario'],
        ['Dispositivo', 'PC de escritorio o laptop', 'PC con pantalla 15" o mayor para uso diario en oficina'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 5.5, 7.5])

    body(doc,
         'Para acceder al sistema desde un dispositivo móvil (celular o tablet), '
         'se recomienda usar la URL de la aplicación móvil ciudadana, que está optimizada '
         'para pantallas pequeñas. El portal administrativo completo está optimizado '
         'para pantallas de escritorio.')

    note_box(doc,
             'Internet Explorer no es compatible con la plataforma bajo ninguna versión. '
             'Si el equipo del usuario tiene configurado Internet Explorer como navegador '
             'predeterminado, debe instalar Google Chrome o Microsoft Edge antes de acceder.',
             label='Compatibilidad', bcolor='B77910', bgcolor='FFFBEB')

    heading(doc, '1.8. Proceso de Inicio de Sesión', level=2)

    body(doc,
         'El proceso de autenticación en la plataforma sigue los siguientes pasos:')

    numbered(doc, 'Abrir el navegador web e ingresar la URL: https://geoportal.munialgarrobo.cl')
    numbered(doc, 'El sistema mostrará la pantalla de login con los campos "Correo electrónico" y "Contraseña".')
    numbered(doc, 'Ingresar el correo electrónico institucional proporcionado por Informática (ejemplo: jperez@munialgarrobo.cl).')
    numbered(doc, 'Ingresar la contraseña asignada. La contraseña distingue mayúsculas de minúsculas.')
    numbered(doc, 'Hacer clic en el botón "Iniciar Sesión".')
    numbered(doc, 'Si las credenciales son correctas, el sistema generará un token de sesión (JWT) y redirigirá al dashboard correspondiente al rol del usuario.')
    numbered(doc, 'Si las credenciales son incorrectas, el sistema mostrará el mensaje "Credenciales inválidas" sin especificar cuál campo es incorrecto (por seguridad). Después de 5 intentos fallidos consecutivos, el sistema bloqueará la IP por 1 minuto.')

    note_box(doc,
             'La sesión tiene una duración de 24 horas. Transcurrido ese tiempo, el sistema '
             'solicitará nuevamente las credenciales. Para cerrar sesión antes del vencimiento, '
             'hacer clic en el nombre de usuario en la esquina superior derecha y seleccionar '
             '"Cerrar Sesión".',
             label='Duración de Sesión', bcolor='2B6CB0', bgcolor='EBF4FF')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. GUÍA DE USO ADMINISTRADOR
# ═══════════════════════════════════════════════════════════════════════════════

def section_2(doc):
    heading(doc, '2. Guía de Uso — Administrador del Sistema')

    body(doc,
         'Esta guía está dirigida exclusivamente al usuario con rol Admin General (nivel 10), '
         'quien tiene acceso completo e irrestricto a todos los módulos y funcionalidades '
         'de la plataforma. El administrador es el responsable máximo de la configuración '
         'del sistema, la gestión de cuentas de usuarios, la supervisión de la actividad, '
         'la ejecución de auditorías de calidad, la gestión del almacenamiento y los '
         'procedimientos de respaldo y migración de datos.')

    note_box(doc,
             'Solo debe existir un número mínimo y necesario de usuarios con rol Admin General '
             '(nivel 10). Se recomienda no más de dos personas con este nivel de acceso: '
             'el encargado de Informática y un respaldo designado. El rol debe ser revocado '
             'inmediatamente cuando el funcionario deje el cargo o cambie de responsabilidades.',
             label='Principio de Menor Privilegio', bcolor='C53030', bgcolor='FFF5F5')

    heading(doc, '2.1. Panel de Administración — Navegación General', level=2)

    body(doc,
         'Al iniciar sesión con una cuenta de nivel 10, el sistema redirige automáticamente '
         'al Panel de Administración. La interfaz está organizada en una barra de navegación '
         'lateral izquierda con acceso a todos los módulos del sistema. '
         'Los módulos exclusivos del administrador aparecen destacados en la navegación.')

    body(doc, 'Estructura del menú de navegación para el Admin General:')

    headers = ['Sección del Menú', 'Submódulo', 'Función']
    rows = [
        ['Dashboard',           '—',                    'Resumen ejecutivo: conteo de proyectos activos, licitaciones en curso, alertas pendientes y actividad reciente'],
        ['Administración',      'Usuarios',              'Gestión completa de cuentas: crear, editar, activar/desactivar usuarios del sistema'],
        ['Administración',      'Divisiones',            'Gestión de las unidades organizacionales de la municipalidad (departamentos, direcciones)'],
        ['Administración',      'Mantenedores',          'Edición de todos los catálogos del sistema: áreas, etapas, estados, financiamientos, sectores, lineamientos'],
        ['SECPLAN',             'Proyectos',             'Acceso total a todos los proyectos de todos los usuarios'],
        ['SECPLAN',             'Licitaciones',          'Gestión y supervisión de todos los procesos licitatorios activos e históricos'],
        ['Geoportal',           '—',                    'Mapa interactivo con acceso a todas las capas temáticas incluyendo las exclusivas del administrador'],
        ['Auditoría',           '—',                    'Motor de análisis de calidad: lanzar auditorías, ver reportes generados, descargar PDF'],
        ['Control de Actividad','—',                    'Historial completo de todas las acciones de todos los usuarios del sistema'],
        ['Gestión de Volumen',  'Exportar / Importar',   'Respaldo y restauración del almacenamiento de archivos del sistema'],
        ['Mi Perfil',           '—',                    'Cambio de contraseña y datos del propio usuario administrador'],
    ]
    make_table(doc, headers, rows, col_widths=[3.8, 3.0, 9.7])

    heading(doc, '2.2. Gestión de Usuarios — Procedimientos Detallados', level=2)

    heading(doc, '2.2.1. Crear un Nuevo Usuario', level=3)

    body(doc,
         'Antes de crear un usuario, el administrador debe tener a mano la siguiente '
         'información: nombre completo del funcionario, correo electrónico institucional, '
         'división o unidad organizacional a la que pertenece, y el nivel de acceso '
         'que corresponde a sus responsabilidades.')

    numbered(doc, 'Ingresar al Panel de Administración → sección "Administración" → subsección "Usuarios".')
    numbered(doc, 'La pantalla mostrará el listado completo de usuarios del sistema con su estado (activo/inactivo) y último acceso.')
    numbered(doc, 'Hacer clic en el botón "Nuevo Usuario" ubicado en la esquina superior derecha del listado.')
    numbered(doc, 'Completar el formulario de creación con todos los campos indicados en la tabla siguiente:')

    headers = ['Campo del Formulario', 'Descripción Detallada', 'Validaciones del Sistema', 'Obligatorio']
    rows = [
        ['Nombre completo',
         'Nombre y apellido(s) del funcionario tal como aparecen en su contrato o credencial municipal',
         'Mínimo 3 caracteres. No acepta solo números.',
         'Sí'],
        ['Correo electrónico',
         'Email institucional del funcionario. Será el identificador único de inicio de sesión. No puede repetirse.',
         'Formato email válido (usuario@dominio.cl). Se valida unicidad en la base de datos.',
         'Sí'],
        ['Contraseña inicial',
         'Contraseña que el usuario usará en su primer acceso. Se recomienda una contraseña temporal que el usuario cambie en su primer login.',
         'Mínimo 8 caracteres. Se almacena cifrada con Bcrypt, nunca en texto plano.',
         'Sí'],
        ['Confirmar contraseña',
         'Repetición de la contraseña para verificar que fue ingresada correctamente',
         'Debe coincidir exactamente con el campo "Contraseña inicial".',
         'Sí'],
        ['Nivel de acceso (Rol)',
         'Nivel jerárquico del usuario según su responsabilidad institucional. Ver tabla de roles en sección 1.5.',
         'Número entero entre 10 y 16. No puede asignarse un rol superior al propio del administrador.',
         'Sí'],
        ['División',
         'Unidad organizacional a la que pertenece el funcionario (ej: SECPLAN, Dirección de Obras, Finanzas)',
         'Debe seleccionar una división existente en el catálogo. Si no existe, crearla primero en "Divisiones".',
         'Sí'],
        ['Estado de la cuenta',
         'Activo: el usuario puede iniciar sesión. Inactivo: la cuenta existe pero el usuario no puede acceder.',
         'Por defecto el valor es "Activo" al crear. Se puede cambiar en cualquier momento.',
         'Sí'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 5.5, 4.5, 2.0])

    numbered(doc, 'Revisar todos los campos antes de confirmar.')
    numbered(doc, 'Hacer clic en "Guardar Usuario". El sistema registrará automáticamente la fecha y hora de creación y el administrador que la realizó.')
    numbered(doc, 'El sistema mostrará un mensaje de confirmación con el ID asignado al nuevo usuario.')
    numbered(doc, 'Comunicar al funcionario su correo de acceso y contraseña inicial de forma segura (preferentemente en persona o por canal cifrado). Solicitar que cambie la contraseña en su primer acceso.')

    note_box(doc,
             'Los usuarios no se eliminan físicamente del sistema para preservar la integridad '
             'de la trazabilidad histórica. Aunque un funcionario deje la municipalidad, sus '
             'datos históricos (proyectos creados, modificaciones realizadas) deben conservarse. '
             'En su lugar, se desactiva la cuenta (estado = Inactivo). Un usuario inactivo no puede '
             'iniciar sesión, pero toda su actividad histórica permanece disponible en los registros.',
             label='Política de Eliminación de Usuarios', bcolor='B77910', bgcolor='FFFBEB')

    heading(doc, '2.2.2. Editar un Usuario Existente', level=3)

    numbered(doc, 'Ingresar al listado de usuarios (Panel de Administración → Usuarios).')
    numbered(doc, 'Localizar al usuario usando la barra de búsqueda (por nombre o correo) o desplazándose por el listado.')
    numbered(doc, 'Hacer clic en el ícono de edición (lápiz) en la fila del usuario.')
    numbered(doc, 'Modificar los campos necesarios. Los campos disponibles son los mismos que en el formulario de creación.')
    numbered(doc, 'Para cambiar la contraseña del usuario: marcar la casilla "Cambiar contraseña", ingresar la nueva contraseña y confirmarla.')
    numbered(doc, 'Para desactivar un usuario: cambiar el campo "Estado de la cuenta" a "Inactivo".')
    numbered(doc, 'Hacer clic en "Guardar Cambios". El sistema registrará la modificación en el Control de Actividad con el valor anterior y el nuevo.')

    heading(doc, '2.2.3. Consultar Historial de Acceso de un Usuario', level=3)

    numbered(doc, 'Ingresar al listado de usuarios y hacer clic en el nombre del usuario.')
    numbered(doc, 'En la ficha del usuario, acceder a la pestaña "Historial de Actividad".')
    numbered(doc, 'El sistema mostrará todas las acciones realizadas por ese usuario: login, vistas, ediciones, descargas y errores.')
    numbered(doc, 'Usar los filtros por rango de fechas para acotar el historial.')

    heading(doc, '2.3. Gestión de Divisiones Organizacionales', level=2)

    body(doc,
         'Las divisiones son las unidades organizacionales de la municipalidad. '
         'Cada usuario pertenece a una división. Las divisiones se utilizan para '
         'organizar los reportes y facilitar la supervisión por unidad de trabajo.')

    body(doc, 'Para crear o editar una división:')
    numbered(doc, 'Ingresar a Panel de Administración → Administración → Divisiones.')
    numbered(doc, 'El listado mostrará todas las divisiones existentes con su nombre, descripción y cantidad de usuarios asignados.')
    numbered(doc, 'Para crear: hacer clic en "Nueva División", ingresar el nombre oficial de la unidad y una descripción, luego guardar.')
    numbered(doc, 'Para editar: hacer clic en el ícono de edición de la división y modificar los campos necesarios.')

    note_box(doc,
             'No se puede eliminar una división si tiene usuarios asignados. Primero debe '
             'reasignar todos los usuarios de esa división a otra, y luego proceder a eliminarla.',
             label='Restricción', bcolor='C53030', bgcolor='FFF5F5')

    heading(doc, '2.4. Gestión de Catálogos (Mantenedores)', level=2)

    body(doc,
         'Los catálogos son las tablas de valores de referencia que se utilizan en los '
         'formularios de proyectos. Cuando un usuario crea o edita un proyecto, los '
         'selectores (desplegables) de campos como "Área Temática", "Etapa del Proyecto" '
         'o "Fuente de Financiamiento" se alimentan de estos catálogos. '
         'El administrador puede agregar, editar o desactivar valores de catálogo '
         'desde el Panel de Administración → Mantenedores.')

    headers = ['Catálogo', 'Descripción del Uso', 'Ejemplos de Valores', 'Campo en Proyecto']
    rows = [
        ['Áreas Temáticas',
         'Clasificación sectorial del proyecto según el tipo de obra o servicio que representa',
         'Vialidad, Saneamiento, Espacios Públicos, Equipamiento, Medioambiente, Educación, Salud',
         'Área Temática'],
        ['Etapas de Proyecto',
         'Fase del ciclo de vida del proyecto en que se encuentra actualmente',
         'Idea, Perfil, Prefactibilidad, Diseño, Licitación, Ejecución, Recepción, Cierre',
         'Etapa del Proyecto'],
        ['Estados de Proyecto',
         'Situación operativa actual del proyecto respecto a su avance y viabilidad',
         'En preparación, En diseño, En licitación, En ejecución, Ejecutado, Paralizado, Descartado',
         'Estado del Proyecto'],
        ['Estados de Postulación',
         'Estado del proceso de postulación del proyecto a la fuente de financiamiento correspondiente',
         'No postulado, Postulado, Con RS aprobado, Con RS rechazado, Aprobado para ejecución',
         'Estado de Postulación'],
        ['Fuentes de Financiamiento',
         'Origen de los recursos económicos con que se financiará o financió el proyecto',
         'FNDR, FRIL, SUBDERE, MINVU, MINSAL, Municipal, Mixto FNDR/Municipal, BIP',
         'Fuente de Financiamiento'],
        ['Sectores Territoriales',
         'División geográfica de la comuna para organización territorial de los proyectos',
         'Centro, Norte, Sur, Rural, Litoral, Sector Las Dichas, Mirasol, Lo Gallardo',
         'Sector Territorial'],
        ['Lineamientos Estratégicos',
         'Eje del Plan Comunal de Desarrollo (PLADECO) al que se adscribe el proyecto',
         'Desarrollo urbano sustentable, Infraestructura y conectividad, Medioambiente, Desarrollo social, Economía local',
         'Lineamiento Estratégico'],
    ]
    make_table(doc, headers, rows, col_widths=[3.2, 4.8, 4.5, 4.0])

    body(doc, 'Procedimiento para agregar un nuevo valor a un catálogo:')
    numbered(doc, 'Ingresar a Panel de Administración → Administración → Mantenedores.')
    numbered(doc, 'Seleccionar el catálogo al que se desea agregar el valor (ej: "Áreas Temáticas").')
    numbered(doc, 'El sistema mostrará los valores existentes con su estado (activo/inactivo).')
    numbered(doc, 'Hacer clic en "Nuevo Valor".')
    numbered(doc, 'Ingresar el nombre del nuevo valor y una descripción opcional.')
    numbered(doc, 'Hacer clic en "Guardar". El nuevo valor estará disponible de inmediato en todos los formularios del sistema.')

    note_box(doc,
             'No se deben eliminar valores de catálogo que estén siendo usados en proyectos '
             'existentes, ya que generaría inconsistencias en los datos. En su lugar, '
             'desactivar el valor (estado = Inactivo): dejará de aparecer en los formularios '
             'nuevos pero se conservará en los proyectos existentes que ya lo tenían asignado.',
             label='Política de Catálogos', bcolor='B77910', bgcolor='FFFBEB')

    heading(doc, '2.5. Módulo de Auditoría Integral — Guía Completa', level=2)

    body(doc,
         'La Auditoría Integral es el motor de análisis automatizado de calidad de la plataforma. '
         'Evalúa la completitud y coherencia de cada proyecto activo en 7 dimensiones con '
         '15 validaciones específicas (V001-V015), genera un puntaje de calidad por dimensión '
         'y un puntaje global, y produce reportes PDF individuales por proyecto con un plan '
         'de acción correctivo priorizado por nivel de urgencia.')

    heading(doc, '2.5.1. Las 7 Dimensiones y sus Validaciones', level=3)

    body(doc,
         'A continuación se describen en detalle las 7 dimensiones evaluadas por el '
         'motor de auditoría, con sus validaciones específicas y criterios de aprobación:')

    headers = ['Dim.', 'Nombre', 'Validaciones Incluidas', 'Peso en Puntaje Global']
    rows = [
        ['D1', 'Identificación y Clasificación',
         'V001: N° de registro presente y con formato válido.\nV002: Área temática asignada desde el catálogo.',
         '10%'],
        ['D2', 'Priorización y Financiamiento',
         'V003: Monto presupuestado > 0.\nV004: Fuente de financiamiento asignada.\nV005: Años de elaboración y ejecución definidos y coherentes (ejecución ≥ elaboración).',
         '15%'],
        ['D3', 'Variables Técnicas',
         'V006: Topografía% ≥ umbral mínimo según etapa.\nV007: Planimetrías% ≥ umbral mínimo según etapa.\nV008: Ingeniería% ≥ umbral mínimo según etapa.\nV009: Perfil técnico-económico% ≥ umbral.\nV010: Documentación% ≥ umbral.',
         '20%'],
        ['D4', 'Estado y Ciclo de Vida',
         'V011: Etapa y estado coherentes entre sí (ej: no puede estar en "Ejecución" con etapa "Idea").\nV012: Estado de postulación asignado.',
         '25%'],
        ['D5', 'Equipo Profesional',
         'V013: Al menos un profesional responsable asignado.',
         '10%'],
        ['D6', 'Aprobaciones y Permisos',
         'V014: Campos de aprobación DOM y SERVIU definidos según corresponda al tipo de proyecto.',
         '10%'],
        ['D7', 'Variables Geográficas',
         'V015: Coordenadas de latitud y longitud presentes y dentro del rango geográfico de la Región de Valparaíso.',
         '10%'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.5, 8.5, 2.7])

    heading(doc, '2.5.2. Niveles de Alerta en los Reportes', level=3)

    body(doc,
         'Cada validación fallida genera una alerta clasificada en cuatro niveles de prioridad. '
         'El reporte PDF incluye una tabla con todas las alertas del proyecto, ordenadas '
         'de mayor a menor urgencia, y el plan de acción recomendado para cada una:')

    bullet(doc,
           'El proyecto presenta una inconsistencia de datos que podría comprometer su postulación, ejecución o validez legal. '
           'Ejemplo: monto = 0 en un proyecto en etapa de Licitación, o etapa y estado incompatibles. '
           'Plazo de subsanación recomendado: inmediato (antes del próximo hito del proyecto).',
           bold_prefix='CRÍTICO', color=RED_CRIT)
    bullet(doc,
           'El proyecto tiene un dato incompleto que representa un riesgo significativo para la gestión. '
           'Ejemplo: sin profesional responsable asignado, sin fuente de financiamiento. '
           'Plazo de subsanación recomendado: 7 días hábiles.',
           bold_prefix='ALTO', color=ORANGE_ALT)
    bullet(doc,
           'El proyecto tiene un dato deseable que mejora la completitud del expediente. '
           'Ejemplo: variables técnicas por debajo del umbral para la etapa actual. '
           'Plazo de subsanación recomendado: 30 días.',
           bold_prefix='MEDIO', color=YELLOW_W)
    bullet(doc,
           'Recomendación de buenas prácticas que mejora la calidad del registro sin ser urgente. '
           'Ejemplo: sin coordenadas geográficas para un proyecto de vialidad. '
           'Plazo de subsanación recomendado: 90 días.',
           bold_prefix='BAJO', color=BLUE_INFO)

    heading(doc, '2.5.3. Procedimiento para Ejecutar una Auditoría', level=3)

    numbered(doc, 'Ingresar al Panel de Administración → sección "Auditoría".')
    numbered(doc, 'En la pantalla principal de Auditoría se mostrarán las auditorías previas realizadas con su fecha, filtros aplicados y cantidad de proyectos evaluados.')
    numbered(doc, 'Hacer clic en "Nueva Auditoría" para configurar y lanzar un nuevo análisis.')
    numbered(doc, 'Aplicar los filtros deseados para determinar el alcance de la auditoría. Los filtros disponibles son:')
    bullet(doc, 'Por área temática (ej: auditar solo proyectos de Vialidad)', level=2)
    bullet(doc, 'Por etapa del proyecto (ej: solo proyectos en etapa de Licitación o Ejecución)', level=2)
    bullet(doc, 'Por estado del proyecto', level=2)
    bullet(doc, 'Por profesional responsable', level=2)
    bullet(doc, 'Sin filtros: evalúa la totalidad de proyectos activos del sistema', level=2)
    numbered(doc, 'Hacer clic en "Lanzar Auditoría". El proceso se ejecuta en segundo plano en el servidor.')
    numbered(doc, 'Monitorear el avance en la barra de progreso que indica el porcentaje de proyectos procesados y el número de proyectos evaluados hasta el momento.')
    numbered(doc, 'Al finalizar (puede tardar entre 30 segundos y varios minutos según la cantidad de proyectos), el sistema enviará una notificación y mostrará el resumen de resultados.')
    numbered(doc, 'Hacer clic en "Ver Reportes" para acceder a los PDF generados. Cada proyecto tiene su propio reporte descargable.')
    numbered(doc, 'Revisar el cuadro de resumen que muestra la distribución de proyectos por puntaje y la cantidad de alertas por nivel (CRÍTICO, ALTO, MEDIO, BAJO).')

    note_box(doc,
             'Los reportes de auditoría se almacenan en /data/auditoria_reportes/ dentro del '
             'volumen persistente del servidor. Cada reporte tiene nombre único con fecha y '
             'código del proyecto. Los reportes persisten entre reinicios del servidor. '
             'Se recomienda lanzar auditorías periódicas (mensualmente mínimo) para mantener '
             'la calidad del registro de proyectos.',
             label='Almacenamiento de Reportes', bcolor='276B3D', bgcolor='EBF5F0')

    heading(doc, '2.6. Módulo de Control de Actividad', level=2)

    body(doc,
         'El Control de Actividad es el registro inmutable de todas las acciones realizadas '
         'en la plataforma. Funciona como bitácora de auditoría institucional: '
         'cada acción queda registrada con quién la realizó, cuándo, desde dónde '
         'y qué datos fueron afectados. Este módulo es de solo lectura para el administrador: '
         'ningún usuario, incluyendo el Admin General, puede modificar o eliminar registros.')

    body(doc, 'Información registrada por cada evento en el sistema:')

    headers = ['Campo del Registro', 'Descripción', 'Ejemplo']
    rows = [
        ['Fecha y hora',       'Timestamp preciso del evento con zona horaria del servidor',       '2026-04-15 14:32:07 UTC-3'],
        ['Usuario',            'Nombre completo y correo del usuario que realizó la acción',       'Juan Pérez (jperez@munialgarrobo.cl)'],
        ['Tipo de acción',     'Categoría estandarizada del evento registrado',                    'LOGIN, VER, CREAR, EDITAR, ELIMINAR, DESCARGAR, ERROR'],
        ['Módulo',             'Sección del sistema donde ocurrió la acción',                      'SECPLAN, LICITACIONES, ADMINISTRACIÓN, AUDITORÍA'],
        ['Entidad afectada',   'Tipo y ID del objeto sobre el que se actuó',                       'Proyecto ID: 142 — "Mejoramiento Plaza de Armas"'],
        ['Dirección IP',       'IP de origen de la petición desde el navegador del usuario',       '192.168.1.55 (red interna municipal)'],
        ['Navegador',          'User-Agent del navegador utilizado en el acceso',                  'Chrome 124 / Windows 10'],
        ['Estado de la acción','Si la operación fue exitosa o generó un error',                    'EXITOSO / ERROR (con código de error)'],
        ['Datos anteriores',   'JSON con los valores del registro antes de la modificación',       '{"estado": "En preparación", "etapa": "Perfil"}'],
        ['Datos nuevos',       'JSON con los valores del registro después de la modificación',     '{"estado": "En licitación", "etapa": "Licitación"}'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 5.5, 7.5])

    body(doc, 'Para consultar el Control de Actividad:')
    numbered(doc, 'Ingresar al Panel de Administración → "Control de Actividad".')
    numbered(doc, 'El sistema mostrará el historial completo en orden cronológico inverso (más reciente primero).')
    numbered(doc, 'Aplicar los filtros disponibles para acotar la búsqueda:')
    bullet(doc, 'Por usuario: seleccionar del desplegable de usuarios del sistema', level=2)
    bullet(doc, 'Por tipo de acción: LOGIN, CREAR, EDITAR, ELIMINAR, DESCARGAR, ERROR', level=2)
    bullet(doc, 'Por módulo: SECPLAN, LICITACIONES, ADMINISTRACIÓN, AUDITORÍA, GEOPORTAL', level=2)
    bullet(doc, 'Por rango de fechas: selector de fecha inicio y fecha fin', level=2)
    bullet(doc, 'Por entidad: buscar por ID o nombre de proyecto/licitación específica', level=2)
    numbered(doc, 'Para ver el detalle completo de una acción (incluyendo datos anteriores y nuevos), hacer clic en la fila del evento.')
    numbered(doc, 'Para exportar el historial filtrado a Excel, hacer clic en "Exportar CSV".')

    note_box(doc,
             'El módulo de Control de Actividad es la herramienta principal para investigar '
             'incidentes: modificaciones no autorizadas, accesos sospechosos o errores reportados '
             'por usuarios. Ante cualquier incidente de seguridad, la primera acción debe ser '
             'revisar el historial de actividad del período afectado.',
             label='Uso en Investigación de Incidentes', bcolor='1A2E4A', bgcolor='EEF2F7')

    heading(doc, '2.7. Monitoreo de Salud del Sistema', level=2)

    body(doc,
         'El administrador puede verificar el estado operacional del servidor en cualquier '
         'momento accediendo al endpoint de salud del sistema. Este endpoint retorna '
         'información sobre el estado de la base de datos, el almacenamiento y el servidor.')

    body(doc,
         'Para consultar el estado del sistema, abrir el navegador e ingresar la siguiente URL:')

    note_box(doc,
             'https://geoportal.munialgarrobo.cl/health\n\n'
             'El sistema retornará un JSON con los siguientes datos:\n'
             '- status: "healthy" (saludable) o "unhealthy" (con problemas)\n'
             '- database.status: "connected" (BD funcionando) o "error" (BD con problemas)\n'
             '- storage.active_root: ruta del almacenamiento activo (/data/docs o local)\n'
             '- timestamp: fecha y hora exacta de la consulta\n\n'
             'Si el status retorna "unhealthy", contactar inmediatamente a la Unidad de Informática.',
             label='Endpoint de Salud', bcolor='276B3D', bgcolor='EBF5F0')

    heading(doc, '2.8. Procedimiento de Cambio de Contraseña de Usuario', level=2)

    body(doc,
         'Cuando un funcionario olvida su contraseña o requiere un restablecimiento, '
         'el administrador debe seguir este procedimiento:')

    numbered(doc, 'Ingresar al Panel de Administración → Administración → Usuarios.')
    numbered(doc, 'Localizar al usuario por nombre o correo electrónico en la barra de búsqueda.')
    numbered(doc, 'Hacer clic en el ícono de edición (lápiz) en la fila del usuario.')
    numbered(doc, 'En el formulario de edición, marcar la casilla "Cambiar contraseña".')
    numbered(doc, 'Ingresar una nueva contraseña temporal en el campo "Nueva contraseña". Se recomienda una combinación de al menos 8 caracteres con letras y números.')
    numbered(doc, 'Confirmar la contraseña en el campo "Confirmar nueva contraseña".')
    numbered(doc, 'Hacer clic en "Guardar Cambios".')
    numbered(doc, 'Comunicar la nueva contraseña temporal al funcionario por canal seguro.')
    numbered(doc, 'Solicitar al funcionario que cambie la contraseña temporal en su próximo acceso desde la sección "Mi Perfil".')

    note_box(doc,
             'El administrador puede cambiar la contraseña de cualquier usuario del sistema. '
             'Sin embargo, el sistema registra este cambio en el Control de Actividad, '
             'indicando qué administrador realizó el restablecimiento y en qué fecha/hora. '
             'Esto garantiza la trazabilidad completa de los cambios de credenciales.',
             label='Trazabilidad de Cambios de Contraseña', bcolor='2B6CB0', bgcolor='EBF4FF')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. GUÍA DE USO PROYECTOS
# ═══════════════════════════════════════════════════════════════════════════════

def section_3(doc):
    heading(doc, '3. Guía de Uso — Módulo Proyectos (SECPLAN)')

    body(doc,
         'El módulo de Proyectos es el núcleo operativo de la plataforma y la razón '
         'de ser de la SECPLAN en el sistema. Permite registrar, gestionar y dar seguimiento '
         'a la totalidad de los proyectos de inversión pública de la municipalidad. '
         'La ficha de proyecto es el expediente digital completo de cada iniciativa: '
         'contiene más de 60 campos organizados en 6 secciones, más subsistemas de '
         'documentos, hitos, observaciones, geomapas y chat asistente IA.')

    body(doc,
         'Acceden a este módulo los roles: Admin General (10), Admin Proyectos (11), '
         'Director de Obras (12), Profesional Técnico (13) y Finanzas (14), '
         'cada uno con permisos diferenciados de lectura y escritura según su rol.')

    heading(doc, '3.1. Dashboard de Proyectos — Pantalla Principal', level=2)

    body(doc,
         'Al ingresar al módulo SECPLAN, el usuario visualiza el tablero principal '
         'que organiza la información de los proyectos en las siguientes áreas:')

    bullet(doc,
           'Cuatro tarjetas de resumen en la parte superior con contadores en tiempo real: '
           'Total de Proyectos, Proyectos en Ejecución, Proyectos con Alertas y Proyectos '
           'Ejecutados este año.',
           bold_prefix='Resumen Estadístico')
    bullet(doc,
           'Tabla paginada con todos los proyectos accesibles para el usuario según su rol. '
           'Cada fila muestra: N° de registro, nombre del proyecto, área temática, etapa, '
           'estado, monto presupuestado, profesional responsable y porcentaje de avance.',
           bold_prefix='Listado de Proyectos')
    bullet(doc,
           'Barra de búsqueda con filtros avanzados combinables: nombre del proyecto, '
           'N° de registro, área temática, etapa, estado, fuente de financiamiento, '
           'profesional responsable, rango de años de ejecución y sector territorial.',
           bold_prefix='Filtros Avanzados')
    bullet(doc,
           'Barra de progreso visual por proyecto que indica el porcentaje de completitud '
           'de los campos de la ficha según el motor de auditoría.',
           bold_prefix='Indicadores de Completitud')
    bullet(doc,
           'Botón de exportación del listado a Excel con los campos visibles en la tabla.',
           bold_prefix='Exportar a Excel')

    heading(doc, '3.2. Crear una Nueva Ficha de Proyecto', level=2)

    body(doc,
         'La ficha de proyecto es el expediente digital centralizado de cada iniciativa '
         'de inversión. Debe ser completada con el máximo detalle posible, ya que la '
         'completitud de sus campos determina directamente el puntaje que obtendrá '
         'en la Auditoría Integral. A continuación se describe el procedimiento de '
         'creación y cada campo del formulario.')

    numbered(doc, 'Hacer clic en el botón "Nuevo Proyecto" en el dashboard del módulo SECPLAN.')
    numbered(doc, 'El sistema abrirá el formulario de creación organizado en 6 secciones. Completar cada sección según se indica a continuación.')

    heading(doc, '3.2.1. Sección A — Identificación y Clasificación', level=3)

    body(doc,
         'Esta sección contiene los datos de identificación del proyecto. '
         'Son los campos más básicos y todos son obligatorios para guardar la ficha.')

    headers = ['Campo', 'Descripción Detallada', 'Validación / Formato', 'Oblig.']
    rows = [
        ['N° de Registro',
         'Código único del proyecto asignado por SECPLAN. Identifica al proyecto en toda la plataforma y en los documentos oficiales. Ej: 2024-001, FNDR-2025-017.',
         'Texto libre, máx. 20 caracteres. Debe ser único en el sistema: no pueden existir dos proyectos con el mismo código.',
         'Sí'],
        ['Nombre del Proyecto',
         'Denominación oficial y completa del proyecto tal como aparece en los documentos de postulación y contratos. Debe ser descriptivo y estandarizado.',
         'Texto libre, máx. 200 caracteres. Formato recomendado: [Verbo] + [Objeto] + [Ubicación]. Ej: "Mejoramiento Pasaje Los Aromos, Sector Norte".',
         'Sí'],
        ['Área Temática',
         'Clasificación sectorial del proyecto. Determina en qué categoría aparecerá en el Geoportal y en los reportes estadísticos.',
         'Selector con valores del catálogo "Áreas Temáticas". Solo se pueden seleccionar valores activos.',
         'Sí'],
        ['Lineamiento Estratégico',
         'Eje del Plan Comunal de Desarrollo (PLADECO) al que se adscribe el proyecto. Permite alinear el portafolio con la planificación comunal.',
         'Selector con valores del catálogo "Lineamientos Estratégicos".',
         'Sí'],
        ['Sector Territorial',
         'Zona geográfica de la comuna donde se localiza o tiene impacto el proyecto.',
         'Selector con valores del catálogo "Sectores Territoriales".',
         'Sí'],
        ['Unidad Vecinal',
         'Junta de vecinos, comité o unidad vecinal relacionada directamente con el proyecto.',
         'Texto libre, máx. 100 caracteres. Ej: "Junta de Vecinos Mirasol N°15".',
         'No'],
    ]
    make_table(doc, headers, rows, col_widths=[3.0, 6.0, 5.5, 1.5])

    heading(doc, '3.2.2. Sección B — Información Financiera', level=3)

    headers = ['Campo', 'Descripción Detallada', 'Validación / Formato', 'Oblig.']
    rows = [
        ['Monto Presupuestado',
         'Valor total estimado del proyecto en pesos chilenos (CLP), sin puntos ni comas. Incluye todos los ítems: obras, profesionales, IVA, imprevistos.',
         'Número entero ≥ 0. El sistema formatea automáticamente con separadores de miles.',
         'Sí'],
        ['Fuente de Financiamiento',
         'Origen de los recursos con que se ejecutará el proyecto. Determina el instrumento de postulación aplicable.',
         'Selector con valores del catálogo "Fuentes de Financiamiento".',
         'Sí'],
        ['¿Financiamiento municipal adicional?',
         'Indica si el municipio aporta recursos además del financiamiento externo principal.',
         'Casilla Sí/No.',
         'No'],
        ['Monto financiamiento municipal',
         'Monto del aporte municipal en pesos. Solo visible si el campo anterior está marcado.',
         'Número entero ≥ 0. Solo visible cuando tiene aporte municipal = Sí.',
         'Condicional'],
        ['Año de Elaboración',
         'Año en que se elaboró la propuesta o idea de proyecto.',
         'AAAA. Debe ser ≤ año actual.',
         'Sí'],
        ['Año de Ejecución',
         'Año programado o real en que se ejecutará el proyecto.',
         'AAAA. Debe ser ≥ al año de elaboración.',
         'Sí'],
        ['Fecha de Postulación',
         'Fecha en que se envió la postulación formal a la fuente de financiamiento.',
         'DD/MM/AAAA.',
         'No'],
        ['Estado de Postulación',
         'Estado actual del proceso de postulación al financiamiento.',
         'Selector con valores del catálogo "Estados de Postulación".',
         'Sí'],
    ]
    make_table(doc, headers, rows, col_widths=[3.2, 5.8, 5.2, 1.8])

    heading(doc, '3.2.3. Sección C — Variables Técnicas de Avance', level=3)

    body(doc,
         'Las variables técnicas son porcentajes (0-100%) que representan el grado de completitud '
         'de cada componente técnico del proyecto. Son la base de la Dimensión 3 del motor '
         'de auditoría (peso 20%). Los umbrales mínimos dependen de la etapa del proyecto: '
         'un proyecto en Licitación debe tener Ingeniería ≥ 80% y Planimetrías ≥ 90%.')

    headers = ['Variable', 'Qué Mide Exactamente', 'Se Completa al 100% Cuando...']
    rows = [
        ['Avance Total (%)',           'Porcentaje global de avance considerando todos los componentes técnicos, financieros y administrativos.',      'El proyecto fue recibido definitivamente y todos los componentes están completos.'],
        ['Topografía (%)',             'Completitud del levantamiento topográfico del área de intervención.',                                          'El levantamiento fue ejecutado, revisado y aprobado por el profesional a cargo.'],
        ['Planimetrías (%)',           'Porcentaje de planos y planimetrías disponibles, revisados y aprobados.',                                      'El 100% de los planos del expediente técnico están finalizados y firmados.'],
        ['Ingeniería (%)',             'Completitud de la ingeniería de detalle: cálculos, especificaciones técnicas, memoria de cálculo.',             'Toda la ingeniería de detalle está terminada, revisada y en el expediente técnico.'],
        ['Perfil Técnico-Económico (%)','Completitud del documento de Perfil Técnico-Económico para postulación a fondos.',                           'El Perfil está completo, visado por SECPLAN y listo para postulación.'],
        ['Documentación (%)',          'Porcentaje de documentación general disponible (certificados, autorizaciones, documentos administrativos).',    'Toda la documentación requerida por la fuente de financiamiento está adjunta.'],
    ]
    make_table(doc, headers, rows, col_widths=[3.0, 6.5, 7.0])

    heading(doc, '3.2.4. Sección D — Estado y Ciclo de Vida', level=3)

    headers = ['Campo', 'Descripción Detallada', 'Valores Posibles / Validación']
    rows = [
        ['Etapa del Proyecto',
         'Fase del ciclo de vida en que se encuentra. El motor de auditoría verifica coherencia etapa-estado.',
         'Catálogo: Idea, Perfil, Prefactibilidad, Diseño, Licitación, Ejecución, Recepción, Cierre.'],
        ['Estado del Proyecto',
         'Situación operativa actual. Complementa a la etapa describiendo qué está ocurriendo con el proyecto.',
         'Catálogo: En preparación, En diseño, En licitación, En ejecución, Ejecutado, Paralizado, Descartado.'],
        ['Próximos Pasos',
         'Acciones concretas definidas por el equipo para avanzar el proyecto. Se recomienda: acción + responsable + fecha objetivo.',
         'Texto libre, máx. 500 caracteres.'],
        ['Observaciones Generales',
         'Notas administrativas, técnicas o de contexto que complementan la información del proyecto.',
         'Texto libre, máx. 1000 caracteres.'],
    ]
    make_table(doc, headers, rows, col_widths=[3.2, 6.5, 6.8])

    heading(doc, '3.2.5. Sección E — Equipo Profesional', level=3)

    headers = ['Campo', 'Descripción', 'Validación']
    rows = [
        ['Profesional Responsable 1 a 5',
         'Nombre(s) de los profesionales a cargo. Hasta 5 integrantes. Al menos el campo 1 debe completarse para aprobar validación V013.',
         'Texto libre. Profesional 1 es obligatorio para cumplir la Dimensión 5 de auditoría.'],
        ['Dupla Supervisora Técnica',
         'Nombre del profesional que ejerce supervisión técnica (ej: jefe SECPLAN).',
         'Texto libre.'],
        ['Dupla Supervisora Financiera',
         'Nombre del profesional que ejerce supervisión financiera y presupuestaria.',
         'Texto libre.'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 6.5, 6.5])

    heading(doc, '3.2.6. Sección F — Aprobaciones y Georreferenciación', level=3)

    headers = ['Campo', 'Descripción', 'Formato / Validación']
    rows = [
        ['Aprobación DOM',
         'Indica si el proyecto cuenta con aprobación de la Dirección de Obras Municipales.',
         'Selector: Sí / No / No Aplica.'],
        ['Aprobación SERVIU',
         'Indica si el proyecto cuenta con aprobación del SERVIU regional.',
         'Selector: Sí / No / No Aplica.'],
        ['Latitud',
         'Coordenada geográfica decimal del punto de intervención. Región de Valparaíso: entre -32 y -34.',
         'Decimal 6 decimales. Ej: -33.362841. Sistema valida rango regional.'],
        ['Longitud',
         'Coordenada geográfica decimal. Región de Valparaíso: entre -71 y -72.',
         'Decimal 6 decimales. Ej: -71.672945. Sistema valida rango regional.'],
    ]
    make_table(doc, headers, rows, col_widths=[3.0, 7.0, 6.5])

    note_box(doc,
             'Para obtener coordenadas de un punto en Google Maps: clic derecho sobre '
             'el punto de intervención → "¿Qué hay aquí?". Las coordenadas aparecen '
             'en formato decimal (latitud, longitud) en la parte inferior de la pantalla.',
             label='Cómo Obtener Coordenadas', bcolor='276B3D', bgcolor='EBF5F0')

    numbered(doc, 'Hacer clic en "Guardar Proyecto". El sistema registra automáticamente la fecha de creación y el usuario responsable.')
    numbered(doc, 'El sistema redirige a la ficha completa desde donde se accede a las pestañas: Documentos, Hitos, Observaciones, Geomapa y Chat IA.')

    heading(doc, '3.3. Módulo de Documentos del Proyecto', level=2)

    body(doc,
         'Cada proyecto tiene un repositorio de documentos adjuntos organizado por tipo. '
         'El sistema registra automáticamente quién subió cada documento y cuándo.')

    numbered(doc, 'Abrir la ficha del proyecto → pestaña "Documentos".')
    numbered(doc, 'El sistema muestra el listado de documentos existentes ordenados por fecha de subida.')
    numbered(doc, 'Hacer clic en "Adjuntar Documento".')
    numbered(doc, 'Seleccionar el tipo de documento: Planimetría, Informe Técnico, Decreto, Contrato, Certificado, Permiso, Fotografía u Otro.')
    numbered(doc, 'Escribir una descripción clara del documento (ej: "Plano de plantas arquitectura v2.1 aprobado DOM").')
    numbered(doc, 'Seleccionar el archivo desde el computador (PDF, DOCX, XLSX, JPG, PNG — máx. 50 MB).')
    numbered(doc, 'Hacer clic en "Subir Documento". El sistema valida el tipo real del archivo (magic bytes) antes de aceptarlo.')
    numbered(doc, 'Al finalizar, el documento aparece en el listado con nombre, tipo, tamaño, fecha y usuario.')

    note_box(doc,
             'El sistema valida el tipo real del archivo mediante inspección de los primeros '
             'bytes (magic bytes), no solo la extensión. No acepta ejecutables (.exe, .bat, .sh) '
             'aunque se les cambie la extensión a .pdf.',
             label='Seguridad de Archivos', bcolor='C53030', bgcolor='FFF5F5')

    heading(doc, '3.4. Módulo de Hitos del Proyecto', level=2)

    body(doc,
         'Los hitos registran eventos clave del ciclo de vida del proyecto con fecha y descripción. '
         'Los hitos ingresados aparecen en el calendario integrado del sistema.')

    numbered(doc, 'Abrir la ficha del proyecto → pestaña "Hitos".')
    numbered(doc, 'Hacer clic en "Registrar Nuevo Hito".')
    numbered(doc, 'Seleccionar la categoría del hito (38 categorías predefinidas, ej: Inicio de obra, Recepción provisoria, Postulación a financiamiento).')
    numbered(doc, 'Ingresar la fecha del evento (DD/MM/AAAA).')
    numbered(doc, 'Agregar una descripción detallada: qué ocurrió, quiénes participaron, resultados o compromisos.')
    numbered(doc, 'Hacer clic en "Guardar Hito".')

    heading(doc, '3.5. Módulo de Geomapas del Proyecto', level=2)

    body(doc,
         'Cada proyecto puede tener capas GeoJSON (polígonos, trazas, puntos) que se '
         'visualizan en el Geoportal superpuestas al mapa base.')

    numbered(doc, 'Abrir la ficha del proyecto → pestaña "Geomapa".')
    numbered(doc, 'Hacer clic en "Nueva Capa GeoJSON".')
    numbered(doc, 'Ingresar un nombre descriptivo para la capa (ej: "Traza vial propuesta").')
    numbered(doc, 'Pegar el contenido GeoJSON generado con geojson.io o QGIS.')
    numbered(doc, 'Hacer clic en "Guardar Capa". El sistema valida que el JSON sea GeoJSON válido antes de guardar.')
    numbered(doc, 'La capa estará disponible en el Geoportal de inmediato.')

    heading(doc, '3.6. Asistente IA por Proyecto', level=2)

    body(doc,
         'Cada proyecto tiene integrado un asistente IA (GLM-4.5 Flash) accesible '
         'desde la pestaña "Chat IA" de la ficha. El asistente conoce los datos del '
         'proyecto y puede responder preguntas en lenguaje natural: estado de variables '
         'técnicas, documentos faltantes, resúmenes para reuniones, recomendaciones de avance.')

    note_box(doc,
             'El asistente IA requiere sesión activa y no conserva historial entre sesiones. '
             'Sus respuestas son orientativas y no constituyen resoluciones administrativas.',
             label='Limitaciones del Asistente IA', bcolor='2B6CB0', bgcolor='EBF4FF')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. GUÍA DE USO LICITACIONES
# ═══════════════════════════════════════════════════════════════════════════════

def section_4(doc):
    heading(doc, '4. Guía de Uso — Módulo Licitaciones')

    body(doc,
         'El módulo de Licitaciones gestiona el proceso licitatorio mediante un workflow '
         'de 32 pasos estandarizados, desde el Acuerdo de Concejo Municipal hasta la '
         'Recepción Definitiva de la obra o servicio. Acceden: Admin General (10), '
         'Admin Proyectos (11) y Director de Obras (12).')

    heading(doc, '4.1. Los 32 Pasos del Workflow Licitatorio', level=2)

    body(doc, 'ETAPA 1 — Acuerdo y Aprobación Institucional (Pasos 1–5):', bold=True, color=NAVY_MED)

    headers = ['Paso', 'Nombre', 'Descripción', 'Documentos Típicos']
    rows = [
        ['1', 'Acuerdo de Concejo Municipal',         'El Concejo aprueba en sesión la autorización para licitar la obra o servicio.',                                             'Acta de sesión, Extracto del acuerdo'],
        ['2', 'Decreto de Aprobación de Convenio',    'Si el proyecto es financiado externamente (GORE, SUBDERE), se firma el convenio y se emite el decreto alcaldicio.',         'Convenio, Decreto alcaldicio'],
        ['3', 'Elaboración de Bases Administrativas', 'El equipo redacta las Bases Administrativas y Técnicas que regirán el proceso.',                                            'Borrador de Bases, EETT'],
        ['4', 'Visación por Asesor Jurídico',         'El Asesor Jurídico revisa las Bases para verificar conformidad legal.',                                                     'Bases con VB° Jurídico'],
        ['5', 'Decreto de Aprobación de Bases',       'El Alcalde emite decreto alcaldicio aprobando formalmente las Bases.',                                                       'Decreto Alcaldicio de Bases'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 2 — Publicación en Mercado Público (Pasos 6–9):', bold=True, color=NAVY_MED)
    rows = [
        ['6',  'Publicación en www.mercadopublico.cl', 'Publicación del llamado con Bases, planos y anexos en el portal transaccional del Estado.',   'Ficha de publicación, ID licitación'],
        ['7',  'Período de Preguntas',                 'Los oferentes formulan preguntas por el portal. La entidad debe responder dentro del plazo.',  'Preguntas y respuestas del portal'],
        ['8',  'Circular de Aclaraciones',             'Documento formal que consolida todas las preguntas y respuestas. Se publica en el portal.',    'Circular publicada en Mercado Público'],
        ['9',  'Cierre de Recepción de Ofertas',       'Fecha y hora límite para que oferentes suban sus propuestas. Después no se aceptan nuevas.',   'Confirmación de cierre en portal'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 3 — Evaluación de Ofertas (Pasos 10–14):', bold=True, color=NAVY_MED)
    rows = [
        ['10', 'Apertura de Ofertas Técnicas',         'La comisión abre y revisa las ofertas técnicas, verificando requisitos habilitantes.',         'Acta de apertura técnica'],
        ['11', 'Apertura de Ofertas Económicas',       'Apertura de las propuestas económicas de quienes pasaron la revisión técnica.',                'Acta de apertura económica'],
        ['12', 'Evaluación por Comisión Técnica',      'La comisión evalúa ofertas aplicando criterios y ponderaciones de las Bases.',                 'Cuadro comparativo, actas parciales'],
        ['13', 'Informe Final de Evaluación',          'Consolida resultados con el ranking final de oferentes según criterios de las Bases.',          'Informe Final firmado por la comisión'],
        ['14', 'Acta de Adjudicación',                 'La comisión propone formalmente al Alcalde la adjudicación al mejor oferente.',                 'Acta de Adjudicación firmada'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 4 — Adjudicación y Notificación (Pasos 15–19):', bold=True, color=NAVY_MED)
    rows = [
        ['15', 'Decreto Alcaldicio de Adjudicación',   'El Alcalde adjudica formalmente la licitación al oferente seleccionado.',                      'Decreto Alcaldicio de Adjudicación'],
        ['16', 'Publicación de Resultado',             'Se publica en el portal el resultado indicando el oferente adjudicado y el monto.',            'Publicación en Mercado Público'],
        ['17', 'Notificación a No Adjudicados',        'Se notifica a los oferentes no adjudicados informando el resultado del proceso.',              'Cartas de notificación'],
        ['18', 'Impugnación (si aplica)',              'Período de recursos de impugnación de oferentes no adjudicados. Si no hay, se avanza.',        'Resolución de impugnaciones'],
        ['19', 'Resolución Exenta de Adjudicación',   'Formaliza definitivamente la adjudicación una vez resueltas las impugnaciones.',                'Resolución Exenta'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 5 — Contrato y Garantías (Pasos 20–24):', bold=True, color=NAVY_MED)
    rows = [
        ['20', 'Elaboración del Contrato',             'El equipo jurídico elabora el contrato basado en las Bases y la oferta adjudicada.',           'Borrador del contrato'],
        ['21', 'Firma del Contrato',                   'El Alcalde y el representante legal de la empresa suscriben el contrato.',                      'Contrato firmado por ambas partes'],
        ['22', 'Entrega de Garantía de Fiel Cumplimiento','El contratista entrega la boleta de garantía según lo estipulado en las Bases.',            'Boleta de garantía bancaria / Póliza'],
        ['23', 'Garantía de Correcta Inversión',       'Si hay anticipo, el contratista entrega garantía de correcta inversión del anticipo.',         'Boleta de garantía de anticipos'],
        ['24', 'Decreto de Inicio de Obras',           'El Alcalde autoriza mediante decreto el inicio de las obras o servicios.',                     'Decreto de inicio de obra'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 6 — Ejecución y Control (Pasos 25–29):', bold=True, color=NAVY_MED)
    rows = [
        ['25', 'Inicio Efectivo de Obra',              'Fecha real de inicio físico en terreno. Se levanta acta con ITO y contratista.',              'Acta de inicio de obra'],
        ['26', 'Inspecciones Técnicas (ITO)',           'Visitas periódicas del ITO al terreno para verificar avance y calidad.',                       'Informes de visita ITO, Libro de obras'],
        ['27', 'Estados de Pago',                      'El contratista presenta cobros por avance. El ITO los aprueba para tramitar el pago.',         'Estado de pago N°X, VB° ITO, Decreto'],
        ['28', 'Modificaciones de Contrato',           'Si se requieren obras adicionales o supresiones, se tramita decreto de modificación.',         'Decreto de modificación de contrato'],
        ['29', 'Informe Final de Obra',                'El ITO certifica la terminación total de los trabajos y su conformidad con el contrato.',      'Informe final del ITO'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    body(doc, 'ETAPA 7 — Cierre y Recepción (Pasos 30–32):', bold=True, color=NAVY_MED)
    rows = [
        ['30', 'Recepción Provisoria',                 'La comisión recepciona provisionalmente la obra. Se identifican observaciones pendientes.',    'Acta de Recepción Provisoria, Observaciones'],
        ['31', 'Corrección de Observaciones / Liquidación','El contratista corrige observaciones. Se tramita la liquidación final del contrato.',      'Informe de correcciones, Res. de Liquidación'],
        ['32', 'Recepción Definitiva',                 'Último paso: recepción definitiva con obra terminada. Se devuelven las garantías.',            'Acta de Recepción Definitiva'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 3.8, 7.2, 4.7])

    heading(doc, '4.2. Crear una Nueva Licitación', level=2)

    numbered(doc, 'Ingresar al módulo Licitaciones → hacer clic en "Nueva Licitación".')
    numbered(doc, 'Completar el formulario: proyecto asociado, nombre oficial, ID Mercado Público (si ya existe), monto estimado en CLP y modalidad (LP, LE o TD).')
    numbered(doc, 'Hacer clic en "Crear Licitación". El sistema generará automáticamente los 32 pasos en estado "Pendiente".')

    note_box(doc,
             'Solo puede existir una licitación activa por proyecto simultáneamente. '
             'Intentar crear una segunda licitación activa para el mismo proyecto generará '
             'el error: "Este proyecto ya tiene una licitación activa."',
             label='Restricción por Proyecto', bcolor='C53030', bgcolor='FFF5F5')

    heading(doc, '4.3. Actualizar el Estado de un Paso', level=2)

    numbered(doc, 'Abrir la licitación desde el dashboard → localizar el paso a actualizar.')
    numbered(doc, 'Hacer clic en el ícono de edición del paso.')
    numbered(doc, 'Actualizar: Estado (Pendiente/En proceso/Completado/Atrasado), Fecha Planificada, Fecha Real (al completar), Observaciones y Documentos Adjuntos.')
    numbered(doc, 'Guardar. El cambio queda registrado en el Control de Actividad.')

    body(doc, 'Indicadores visuales del workflow: pasos completados en verde, en proceso en azul, atrasados en rojo, pendientes en gris.')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 5. GUÍA DE USO GEOPORTAL
# ═══════════════════════════════════════════════════════════════════════════════

def section_5(doc):
    heading(doc, '5. Guía de Uso — Geoportal (Mapa Interactivo)')

    body(doc,
         'El Geoportal es el módulo de inteligencia territorial. Visualiza los proyectos '
         'municipales, licitaciones, reportes ciudadanos y capas GIS personalizadas '
         'sobre un mapa interactivo de la comuna. La información es en tiempo real: '
         'cualquier proyecto con coordenadas aparece automáticamente en el mapa. '
         'Acceden todos los roles internos (10-16) con capas diferenciadas por rol.')

    heading(doc, '5.1. Elementos de la Interfaz del Mapa', level=2)

    headers = ['Elemento', 'Ubicación', 'Función']
    rows = [
        ['Mapa principal',          'Área central (80% de la pantalla)', 'Cartografía interactiva: desplazar con clic+arrastrar, zoom con rueda del ratón o botones +/-.'],
        ['Panel de Capas',          'Esquina superior derecha del mapa',  'Control de capas temáticas: activar/desactivar cada capa independientemente.'],
        ['Panel de Filtros',        'Panel lateral izquierdo desplegable','Filtros combinables para acotar proyectos visibles por área, etapa, estado, financiamiento, etc.'],
        ['Barra de búsqueda',       'Parte superior del panel lateral',   'Buscar proyecto por nombre o código. Al seleccionar, el mapa se centra en el marcador.'],
        ['Panel de Información',    'Panel lateral derecho (al clic en marcador)','Ficha resumida del proyecto: código, nombre, área, etapa, estado, monto, avance, profesional.'],
        ['Botón de geolocalización','Esquina inferior derecha',           'Centra el mapa en la ubicación del dispositivo (requiere permiso de geolocalización del navegador).'],
        ['Control de zoom',         'Esquina superior izquierda',         'Botones + y – para ampliar/reducir.'],
        ['Escala',                  'Esquina inferior izquierda',         'Escala cartográfica actual (metros por cm en pantalla).'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 3.5, 9.5])

    heading(doc, '5.2. Capas Temáticas y Roles de Acceso', level=2)

    headers = ['Capa', 'Contenido', 'Geometría', 'Roles']
    rows = [
        ['Proyectos SECPLAN',       'Proyectos con coordenadas registradas. Marcador codificado por color según estado.',           'Puntos',               'Todos (10-16)'],
        ['Licitaciones Activas',    'Proyectos con licitación en curso. Marcador con % de avance del workflow.',                    'Puntos con indicador', 'Roles 10-12'],
        ['GeoJSON de Proyectos',    'Geometrías personalizadas (trazas, polígonos) cargadas desde la ficha de cada proyecto.',     'Puntos, líneas, polígonos', 'Roles 10-14'],
        ['Reportes Ciudadanos',     'Reportes de vecinos geolocalizados. Rojo=pendiente, amarillo=en revisión, verde=resuelto.',   'Puntos con foto',      'Roles 10, 11'],
        ['Cámara Municipal',        'Proyectos estratégicos y zonas de intervención prioritaria de alcaldía.',                     'Puntos y polígonos',   'Solo rol 10'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 6.5, 3.0, 3.5])

    heading(doc, '5.3. Codificación Visual de Marcadores', level=2)

    headers = ['Color', 'Estado del Proyecto', 'Significado Operacional']
    rows = [
        ['Verde oscuro', 'Ejecutado',               'Obra terminada con Recepción Definitiva. Entregada a la comunidad.'],
        ['Verde claro',  'En ejecución activa',      'Obra en terreno con contrato vigente y trabajos en curso.'],
        ['Azul',         'En licitación / Adjudicado','Proceso licitatorio activo o adjudicado pendiente de inicio.'],
        ['Amarillo',     'En diseño / Prefactibilidad','Proyecto con ingeniería o perfil en elaboración.'],
        ['Gris',         'Idea / Perfil inicial',    'Etapa temprana, sin financiamiento asegurado.'],
        ['Naranja',      'Paralizado temporalmente', 'Avance detenido por causas externas.'],
        ['Rojo',         'Con alertas críticas',     'Inconsistencias graves detectadas por auditoría que requieren atención inmediata.'],
    ]
    make_table(doc, headers, rows, col_widths=[2.5, 3.5, 10.5])

    heading(doc, '5.4. Filtros de Visualización Combinados', level=2)

    body(doc,
         'Los filtros son combinables para obtener vistas específicas. '
         'Ejemplo: proyectos de Vialidad + FNDR + En Ejecución + Sector Norte '
         'mostrará solo los marcadores que cumplan TODAS las condiciones simultáneamente.')

    headers = ['Filtro', 'Opciones Disponibles', 'Ejemplo de Uso']
    rows = [
        ['Área Temática',        'Selector múltiple del catálogo',              'Solo proyectos de Vialidad para informe de Dirección de Obras'],
        ['Estado del Proyecto',  'En preparación, En ejecución, Ejecutado, etc.','Ver proyectos en ejecución activa para supervisión en terreno'],
        ['Etapa del Proyecto',   'Idea, Perfil, Diseño, Licitación, Ejecución', 'Proyectos en Licitación para revisión prioritaria del workflow'],
        ['Fuente de Financiamiento','Selector múltiple del catálogo',           'Solo proyectos FNDR para informe al GORE'],
        ['Año de Ejecución',     'Selector de rango año inicio - año fin',      'Proyectos programados para el año en curso'],
        ['Profesional Responsable','Selector con profesionales del sistema',    'Proyectos de un profesional específico para supervisión'],
        ['Sector Territorial',   'Selector múltiple del catálogo',              'Proyectos del Sector Norte para planificación de obras'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 5.0, 8.0])

    heading(doc, '5.5. Atención de Reportes Ciudadanos desde el Mapa', level=2)

    numbered(doc, 'Hacer clic en el marcador de un reporte ciudadano en el mapa.')
    numbered(doc, 'Revisar la descripción, categoría, fotografía y fecha del reporte en el panel lateral.')
    numbered(doc, 'Hacer clic en "Actualizar Estado" para cambiar: Pendiente → En Revisión → En Atención → Resuelto.')
    numbered(doc, 'Agregar un comentario de respuesta indicando la acción tomada.')
    numbered(doc, 'Guardar. El marcador actualiza su color visualmente de inmediato.')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 6. GUÍA DE MIGRACIÓN DEL PORTAL
# ═══════════════════════════════════════════════════════════════════════════════

def section_6(doc):
    heading(doc, '6. Guía de Migración del Portal')

    body(doc,
         'Esta guía describe los procedimientos completos para respaldo, restauración '
         'y migración de datos de la plataforma Algarrobo BASE2. '
         'El sistema almacena datos en dos componentes independientes que deben '
         'respaldarse por separado: (1) el volumen de archivos Docker (/data) con '
         'documentos, reportes y fotos, y (2) la base de datos PostgreSQL con todos '
         'los datos estructurados. Operaciones reservadas para Admin General (nivel 10).')

    note_box(doc,
             'Nunca realice una migración sin respaldar AMBOS componentes: el volumen '
             'de archivos Y la base de datos. Un respaldo parcial resultará en un sistema '
             'inconsistente al restaurar. Toda operación queda registrada en el Control '
             'de Actividad con usuario, fecha y IP.',
             label='Advertencia Crítica', bcolor='C53030', bgcolor='FFF5F5')

    heading(doc, '6.1. Contenido del Volumen de Archivos (/data)', level=2)

    headers = ['Directorio', 'Contenido', 'Tipo de Archivos']
    rows = [
        ['/data/docs/',               'Documentos adjuntos a proyectos y pasos de licitación',      'PDF, DOCX, XLSX, JPG, PNG'],
        ['/data/auditoria_reportes/', 'Reportes PDF del motor de auditoría integral',               'PDF: auditoria_{id}_{fecha}.pdf'],
        ['/data/fotos_reportes/',     'Fotografías de reportes ciudadanos con metadata EXIF',       'JPG, PNG con coordenadas GPS'],
    ]
    make_table(doc, headers, rows, col_widths=[4.0, 7.0, 5.5])

    heading(doc, '6.2. Exportar el Volumen de Archivos (Respaldo)', level=2)

    numbered(doc, 'Ingresar a Panel de Administración → "Gestión de Volumen".')
    numbered(doc, 'Verificar el espacio utilizado y fecha del último respaldo mostrados en la pantalla.')
    numbered(doc, 'Hacer clic en "Exportar Datos (ZIP)".')
    numbered(doc, 'Esperar a que el sistema comprima el volumen (1-15 minutos según tamaño). No cerrar el navegador.')
    numbered(doc, 'Descargar el archivo: backup_algarrobo_AAAAMMDD_HHMM.zip')
    numbered(doc, 'Almacenar en ubicación segura fuera del servidor: NAS institucional, nube autorizada o disco externo.')
    numbered(doc, 'Verificar que el ZIP contiene las tres carpetas: docs/, auditoria_reportes/, fotos_reportes/.')

    note_box(doc,
             'Frecuencia recomendada: mínimo semanal (viernes al cierre). '
             'Conservar los últimos 4 respaldos semanales y uno mensual permanente. '
             'Documentar cada respaldo: fecha, tamaño, ubicación y responsable.',
             label='Política de Respaldo', bcolor='276B3D', bgcolor='EBF5F0')

    heading(doc, '6.3. Respaldo de la Base de Datos PostgreSQL', level=2)

    body(doc,
         'El respaldo de la BD debe realizarse con pg_dump por la Unidad de Informática '
         'o personal técnico con acceso al servidor. No se puede realizar desde el panel web.')

    note_box(doc,
             'COMANDO DE RESPALDO:\n'
             'pg_dump -h [HOST] -p [PUERTO] -U [USUARIO] -d [BD] -F c -b -v -f backup_bd_AAAAMMDD.dump\n\n'
             'Donde: -h = host del servidor PostgreSQL, -p = puerto (ej: 55112),\n'
             '-U = usuario BD (ej: postgres), -d = nombre BD (ej: neondb),\n'
             '-F c = formato custom comprimido, -b = incluir blobs, -f = archivo destino.\n\n'
             'COMANDO DE RESTAURACIÓN:\n'
             'pg_restore -h [HOST] -p [PUERTO] -U [USUARIO] -d [BD] -v backup_bd_AAAAMMDD.dump\n\n'
             'PRECAUCIÓN: pg_restore sobreescribe datos existentes. Verificar que se está '
             'restaurando en el servidor y base de datos correctos antes de ejecutar.',
             label='Comandos pg_dump / pg_restore', bcolor='1A2E4A', bgcolor='EEF2F7')

    heading(doc, '6.4. Restaurar el Volumen desde el Panel Web', level=2)

    note_box(doc,
             'ATENCIÓN: La importación sobreescribe archivos existentes con el mismo nombre. '
             'Realizar un respaldo del estado actual ANTES de importar.',
             label='Advertencia', bcolor='C53030', bgcolor='FFF5F5')

    numbered(doc, 'Panel de Administración → "Gestión de Volumen" → "Importar Datos (ZIP)".')
    numbered(doc, 'Seleccionar el archivo ZIP de respaldo.')
    numbered(doc, 'El sistema valida el ZIP: sin rutas con "..", sin symlinks, solo carpetas permitidas (docs/, auditoria_reportes/, fotos_reportes/), tamaño ≤ 1 GB.')
    numbered(doc, 'Confirmar importación. El sistema extrae los archivos al volumen.')
    numbered(doc, 'Revisar el informe: archivos importados, omitidos y errores.')

    heading(doc, '6.5. Migración Completa entre Servidores', level=2)

    body(doc, 'FASE 1 — Respaldo del Servidor Origen:', bold=True, color=NAVY_MED)
    numbered(doc, 'Notificar a usuarios la ventana de mantenimiento (fuera de horario laboral).')
    numbered(doc, 'Exportar volumen de archivos (sección 6.2) y ejecutar pg_dump de la BD (sección 6.3).')
    numbered(doc, 'Documentar todas las variables de entorno configuradas en Portainer del servidor origen.')
    numbered(doc, 'Detener los contenedores del servidor origen para evitar escrituras durante la migración.')

    body(doc, 'FASE 2 — Preparación del Servidor Destino:', bold=True, color=NAVY_MED)
    numbered(doc, 'Instalar Docker, Portainer y configurar infraestructura según el Manual de Instalación.')
    numbered(doc, 'Crear la base de datos PostgreSQL vacía en el servidor destino.')
    numbered(doc, 'Configurar las variables de entorno en Portainer destino con los mismos valores del origen.')
    numbered(doc, 'Restaurar la BD con pg_restore y verificar la cantidad de registros en tablas principales.')

    body(doc, 'FASE 3 — Migración del Volumen:', bold=True, color=NAVY_MED)
    numbered(doc, 'Desplegar los contenedores en el servidor destino desde Portainer.')
    numbered(doc, 'Importar el ZIP del volumen desde el panel web del nuevo servidor (sección 6.4).')
    numbered(doc, 'Verificar que /data/docs/, /data/auditoria_reportes/ y /data/fotos_reportes/ tienen los archivos correctos.')

    body(doc, 'FASE 4 — Validación:', bold=True, color=NAVY_MED)
    numbered(doc, 'Consultar /health: verificar status="healthy" y database.status="connected".')
    numbered(doc, 'Iniciar sesión, revisar módulo SECPLAN, descargar un documento de proyecto existente.')
    numbered(doc, 'Verificar que los marcadores del Geoportal aparecen correctamente.')
    numbered(doc, 'Ejecutar una auditoría de prueba y verificar que el PDF se genera.')
    numbered(doc, 'Actualizar el DNS del dominio institucional al nuevo servidor y notificar a los usuarios.')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# 7. RESOLUCIÓN DE ANÁLISIS INFORMÁTICA
# ═══════════════════════════════════════════════════════════════════════════════

def section_7(doc):
    heading(doc, '7. Resolución de Análisis — Unidad de Informática')

    body(doc,
         'Este capítulo documenta los criterios técnicos, lineamientos de uso '
         'y resoluciones institucionales adoptadas por la Unidad de Informática '
         'de la I. Municipalidad de Algarrobo respecto a la plataforma Geoportal '
         'Municipal Algarrobo BASE2, producto del análisis de validación técnica '
         'y auditoría de seguridad realizado durante el primer semestre de 2026.')

    heading(doc, '7.1. Estado Operacional del Sistema', level=2)

    headers = ['Módulo / Componente', 'Estado', 'Resultado', 'Observaciones']
    rows = [
        ['Backend Flask API (9 blueprints)', 'Operativo', 'Validado',      'Tiempo de respuesta promedio < 200ms en condiciones normales.'],
        ['Módulo Proyectos (SECPLAN)',        'Operativo', 'Validado',      'CRUD completo, documentos, hitos, observaciones, geomapas y chat IA verificados.'],
        ['Módulo Licitaciones',              'Operativo', 'Validado',      'Workflow 32 pasos, fechas y documentos adjuntos verificados.'],
        ['Auditoría Integral',               'Operativo', 'Validado',      'Motor 7 dimensiones, 15 validaciones V001-V015 y generación PDF confirmados.'],
        ['Geoportal (Leaflet.js)',            'Operativo', 'Validado',      'Capas GeoJSON, filtros combinados y fichas de detalle verificados.'],
        ['Control de Actividad',             'Operativo', 'Validado',      'Registro en tiempo real, filtros y exportación CSV funcionando.'],
        ['Gestión de Volumen (ZIP)',          'Operativo', 'Validado',      'Export/import con validaciones de seguridad (path traversal, symlinks) confirmados.'],
        ['Asistente IA (GLM-4.5 Flash)',      'Operativo', 'Validado',      'Proxy a ZhipuAI con autenticación de sesión requerida funcionando.'],
        ['Autenticación JWT',                'Operativo', 'Validado',      'Login, token, expiración 24h y cierre manual verificados.'],
        ['App Móvil ciudadana',              'En evaluación','Funcional',   'Reporte y geolocalización validados. Pendiente flujo operativo institucional.'],
        ['Almacenamiento S3',                'En evaluación','No impl.',    'Análisis de migración desde volumen Docker a S3-compatible en curso.'],
    ]
    make_table(doc, headers, rows, col_widths=[4.2, 2.3, 2.2, 7.8])

    heading(doc, '7.2. Variables de Entorno — Referencia Completa', level=2)

    body(doc,
         'El sistema se configura exclusivamente mediante variables de entorno en Portainer. '
         'No deben existir credenciales en el código fuente ni en archivos .env en el servidor.')

    headers = ['Variable', 'Descripción', 'Formato / Ejemplo', 'Oblig.']
    rows = [
        ['DATABASE_URL',
         'Cadena completa de conexión PostgreSQL. Incluye usuario, contraseña, host, puerto y nombre BD.',
         'postgresql://user:pass@host:puerto/nombre_bd',
         'Sí'],
        ['JWT_SECRET_KEY',
         'Clave secreta para firmar tokens JWT. Debe ser larga, aleatoria y única. Rotar anualmente.',
         'Min. 64 caracteres aleatorios. Generar: python -c "import secrets; print(secrets.token_hex(32))"',
         'Sí'],
        ['FLASK_ENV',
         'Entorno de ejecución. En producción deshabilita debug y reloader.',
         '"production" en producción, "development" en local',
         'Sí'],
        ['FLASK_DEBUG',
         'Modo debug detallado. NUNCA "true" en producción — expone información sensible.',
         '"false" en producción',
         'Sí'],
        ['ALLOWED_ORIGINS',
         'Dominios permitidos para CORS. En producción, restringir al dominio institucional.',
         'https://geoportal.munialgarrobo.cl (producción) o * (solo desarrollo)',
         'Sí'],
        ['PORT',
         'Puerto de escucha del servidor Gunicorn dentro del contenedor.',
         '8000',
         'No (default: 8000)'],
        ['MAX_UPLOAD_MB',
         'Tamaño máximo por archivo subido en MB.',
         '50 (producción)',
         'No (default: 50)'],
        ['SESSION_EXPIRY_HOURS',
         'Duración en horas del token JWT de sesión.',
         '24',
         'No (default: 24)'],
        ['ZHIPUAI_API_KEY',
         'Clave API de ZhipuAI para el asistente IA. Obtener desde open.bigmodel.cn.',
         'Clave proporcionada por ZhipuAI',
         'Solo si usa IA'],
    ]
    make_table(doc, headers, rows, col_widths=[3.5, 5.5, 5.5, 2.0])

    heading(doc, '7.3. Hallazgos de Seguridad y Estado de Remediación', level=2)

    headers = ['Ref.', 'Hallazgo', 'Impacto', 'Acción Tomada', 'Estado']
    rows = [
        ['3.1', 'Credenciales de BD en código fuente e historial Git.',
         'Acceso no autorizado a producción desde internet.',
         'Rotación de credenciales. Limpieza historial Git (git filter-repo). Migración a Portainer.',
         'Resuelto'],
        ['3.2', 'RBAC incompleto: endpoints sin verificación de nivel de rol.',
         'Escalamiento de privilegios entre roles.',
         'Decoradores de verificación de rol en todos los endpoints sensibles.',
         'Resuelto'],
        ['3.4', 'CORS wildcard (*) en producción.',
         'Ataques CSRF cross-origin.',
         'Restricción a dominio institucional en ALLOWED_ORIGINS.',
         'Resuelto'],
        ['3.5', 'Validación de archivos solo por extensión sin inspección de contenido.',
         'Subida de archivos maliciosos.',
         'Validación por magic bytes antes de aceptar la subida.',
         'Resuelto'],
        ['4.1', 'Condición de carrera en OCR con múltiples usuarios simultáneos.',
         'Corrupción de archivos y errores 500.',
         'Lock de archivo (fcntl) para serializar operaciones OCR.',
         'Resuelto'],
        ['4.3', 'API Key de ZhipuAI expuesta en código fuente del frontend.',
         'Uso no autorizado con costo económico.',
         'Proxy en backend: frontend llama al backend que mantiene la clave en variable de entorno.',
         'Resuelto'],
        ['5.2', 'Doble escritura en tablas de auditoría por bug en módulo.',
         'Duplicados que distorsionaban reportes históricos.',
         'Corrección del bug y limpieza de duplicados históricos.',
         'Resuelto'],
        ['6.1', 'Fallback silencioso a almacenamiento efímero si /data no estaba montado.',
         'Pérdida total de archivos al reiniciar el contenedor.',
         'Warning explícito en logs y validación en /health endpoint.',
         'Resuelto'],
        ['6.2', 'URLs del servidor hardcodeadas en 11 archivos del frontend.',
         'Frontend de producción apuntaba al servidor de desarrollo.',
         'Centralización en variable API_CONFIG.BASE_URL por entorno.',
         'Resuelto'],
        ['6.3', 'Eliminación física de usuarios causaba error 500 por FK constraint.',
         'Corrupción del historial al eliminar usuarios con actividad.',
         'Cambio a borrado lógico (activo = false).',
         'Resuelto'],
        ['4.2', 'Subconsultas correlacionadas en listado de proyectos (N+1 queries).',
         'Degradación con >200 proyectos (>3 segundos de carga).',
         'Plan: triggers de actualización de contadores en tablas relacionadas.',
         'Plan mejora'],
        ['5.1', 'Triggers internos generando eventos duplicados en trazabilidad.',
         'Ruido en historial mezclando acciones sistema/usuario.',
         'Plan: flag de sesión PostgreSQL para distinguir triggers internos.',
         'Plan mejora'],
        ['7.1', 'JWT en localStorage (vulnerable a XSS).',
         'Robo de token en caso de XSS exitoso.',
         'Documentado: migrar a cookies HttpOnly cuando se consolide el dominio definitivo.',
         'Documentado'],
    ]
    make_table(doc, headers, rows, col_widths=[0.8, 4.5, 3.0, 5.0, 2.2])

    heading(doc, '7.4. Lineamientos Técnicos de Administración', level=2)

    heading(doc, '7.4.1. Gestión de Credenciales', level=3)
    numbered(doc, 'Contraseñas mínimo 8 caracteres. Recomendado: 12+ con mayúsculas, minúsculas, números y símbolos.')
    numbered(doc, 'Revisar mensualmente el listado de usuarios y desactivar cuentas inactivas por más de 90 días corridos.')
    numbered(doc, 'Rotar JWT_SECRET_KEY anualmente. La rotación invalida todas las sesiones activas: realizarla fuera del horario laboral con aviso previo.')
    numbered(doc, 'Credenciales de BD y API Keys exclusivamente en Portainer. Nunca en código fuente ni archivos .env en el servidor.')
    numbered(doc, 'Al desvincularse un funcionario con acceso, desactivar la cuenta inmediatamente y rotar credenciales del servidor si tenía acceso a Portainer.')

    heading(doc, '7.4.2. Política de Respaldo', level=3)
    numbered(doc, 'Respaldo semanal de BD (pg_dump) y del volumen (desde panel web), preferentemente los viernes al cierre.')
    numbered(doc, 'Conservar los últimos 4 respaldos semanales y uno mensual permanente por un año.')
    numbered(doc, 'Almacenar respaldos fuera del servidor de producción: NAS institucional, nube o disco externo custodiado.')
    numbered(doc, 'Documentar cada respaldo: fecha, tamaño, ubicación y responsable en un registro simple.')
    numbered(doc, 'Probar la restauración desde respaldo al menos una vez cada seis meses para verificar la validez de los archivos.')

    heading(doc, '7.4.3. Actualizaciones del Sistema', level=3)
    numbered(doc, 'Coordinar actualizaciones con el equipo de desarrollo con mínimo 48 horas de anticipación.')
    numbered(doc, 'Antes de toda actualización, realizar respaldo completo (BD + volumen) sin excepción.')
    numbered(doc, 'Procedimiento estándar: push al repositorio Git → Portainer → "Pull and Redeploy" en el stack de la aplicación.')
    numbered(doc, 'Si la actualización falla: rollback en Portainer editando el stack al tag o commit anterior y haciendo Redeploy.')
    numbered(doc, 'Después de actualizar: consultar /health, verificar login y módulos principales antes de habilitar para usuarios.')

    heading(doc, '7.5. Requerimientos Pendientes al Equipo de Desarrollo', level=2)

    headers = ['Prioridad', 'Requerimiento', 'Justificación', 'Plazo']
    rows = [
        ['Alta',
         'Script SQL de inicialización completa del esquema de base de datos (CREATE TABLE) para nuevas instalaciones.',
         'Sin este script, una nueva instalación requiere acceso al servidor existente para copiar el esquema.',
         'Antes del despliegue en producción'],
        ['Alta',
         'Documentación completa de todas las variables de entorno con descripción, formato y ejemplos.',
         'Las variables no documentadas generan riesgos en actualizaciones y migraciones.',
         'Antes del despliegue en producción'],
        ['Alta',
         'Eliminación de tablas legacy auditoria y auditoria2 del esquema, previa exportación de datos históricos.',
         'Las tablas legacy generan confusión y pueden causar errores en el motor de auditoría actual.',
         'Primer sprint post-producción'],
        ['Media',
         'Triggers de actualización en tablas relacionadas para eliminar subconsultas N+1 en el listado de proyectos.',
         'Con 300+ proyectos el tiempo de carga supera 3 segundos.',
         'Segundo sprint'],
        ['Media',
         'Flag de sesión PostgreSQL (SET LOCAL app.triggered_by_cascade) para evitar duplicados en la trazabilidad.',
         'Los triggers internos contaminan el historial con eventos de sistema.',
         'Segundo sprint'],
        ['Media',
         'Política de archivado trimestral de control_actividad a tabla control_actividad_archivo.',
         'Sin archivado, la tabla crece indefinidamente degradando el rendimiento.',
         'Tercer sprint'],
        ['Baja',
         'Migración de JWT en localStorage a cookies HttpOnly con SameSite=Strict.',
         'Las cookies HttpOnly no son accesibles por JavaScript, eliminando el riesgo de robo por XSS.',
         'Una vez consolidado el dominio municipal definitivo'],
    ]
    make_table(doc, headers, rows, col_widths=[1.8, 5.5, 5.5, 3.7])

    heading(doc, '7.6. Conclusión y Recomendación de Despliegue', level=2)

    body(doc,
         'La plataforma Geoportal Municipal Algarrobo BASE2 ha demostrado ser una '
         'solución técnicamente sólida, funcionalmente completa y arquitectónicamente '
         'bien estructurada para la gestión integral de los proyectos de inversión '
         'pública de la Municipalidad de Algarrobo.')

    body(doc,
         'El proceso de auditoría y validación técnica del primer semestre 2026 '
         'confirmó que el 100% de los hallazgos Bloqueantes y de prioridad Alta '
         '(10 de 13 hallazgos identificados) fueron remediados satisfactoriamente. '
         'Los 3 hallazgos restantes son mejoras de mediano plazo sin riesgo operativo inmediato.')

    body(doc,
         'La arquitectura de 4 capas segregadas (Proxmox → Nginx CT → VM Docker → '
         'PostgreSQL CT) con acceso externo por Túnel Cloudflare sin puertos abiertos '
         'cumple con los estándares de seguridad requeridos para sistemas de información municipal.')

    body(doc,
         'La Unidad de Informática recomienda proceder al despliegue en producción '
         'condicionado a: (a) entrega del script SQL de inicialización y documentación '
         'completa de variables de entorno, y (b) completar las pruebas de aceptación '
         'de usuario (UAT) con el equipo SECPLAN validando todos los flujos principales '
         'con datos reales de la municipalidad.')

    note_box(doc,
             'Documento técnico-administrativo elaborado por la Unidad de Informática '
             'de la I. Municipalidad de Algarrobo en el marco del proceso de implementación '
             'del Geoportal Municipal. Versión 2.0 — Mayo 2026.',
             label='I. Municipalidad de Algarrobo', bcolor='1A2E4A', bgcolor='EEF2F7')


# ═══════════════════════════════════════════════════════════════════════════════
# ENSAMBLAJE Y GUARDADO
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_DARK

    build_cover(doc)
    section_1(doc)
    section_2(doc)
    section_3(doc)
    section_4(doc)
    section_5(doc)
    section_6(doc)
    section_7(doc)

    doc.save(OUTPUT)
    print(f'Documento generado: {OUTPUT}')


if __name__ == '__main__':
    main()

